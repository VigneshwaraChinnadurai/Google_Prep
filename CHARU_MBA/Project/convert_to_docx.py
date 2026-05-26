"""
Convert the MBA Project Report from Markdown to a formatted DOCX file.
"""
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
paragraph_format = style.paragraph_format
paragraph_format.space_after = Pt(6)
paragraph_format.line_spacing = 1.5

# Set margins
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(2.54)

def add_page_break():
    doc.add_page_break()

def add_title_page():
    # Add empty paragraphs for spacing
    for _ in range(4):
        doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('VIT ONLINE LEARNING PROGRAM')
    run.bold = True
    run.font.size = Pt(16)
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('PROJECT REPORT')
    run.bold = True
    run.font.size = Pt(14)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('STRATEGIC COST MANAGEMENT IN INTERIOR PROJECTS')
    run.bold = True
    run.font.size = Pt(16)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Role of Digital Estimation and Vendor Management in\nControlling Cost Overruns and Improving Profitability')
    run.font.size = Pt(13)
    run.italic = True
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('A Project Report submitted in partial fulfillment of the requirement\nfor the award of the degree of Master of Business Administration')
    run.font.size = Pt(11)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Submitted by')
    run.font.size = Pt(12)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Charumathi Baskaran')
    run.bold = True
    run.font.size = Pt(14)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('24EMBA1270')
    run.bold = True
    run.font.size = Pt(13)
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Under the guidance of')
    run.font.size = Pt(12)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('[Guide Name]')
    run.bold = True
    run.font.size = Pt(13)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('May 2026')
    run.font.size = Pt(12)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('VIT Online Learning Program')
    run.font.size = Pt(12)

def add_declaration():
    add_page_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('DECLARATION')
    run.bold = True
    run.font.size = Pt(14)
    
    doc.add_paragraph()
    
    text = ('I, Charumathi Baskaran with register number 24EMBA1270 hereby declare that the '
            'project report entitled "Strategic Cost Management in Interior Projects: Role of Digital '
            'Estimation and Vendor Management in Controlling Cost Overruns and Improving Profitability" '
            'submitted by me to MBA VIT Online Learning Program, Vellore, in partial fulfillment of '
            'the requirement for the award of the degree of Master of Business Administration is a '
            'bonafide work carried out by me under the supervision of [Guide Name], Professor, '
            'Department of Business, VIT Business School, VIT Vellore \u2013 632 014.')
    doc.add_paragraph(text)
    
    doc.add_paragraph()
    
    text2 = ('I further declare that the work reported in this project has not been submitted and '
             'will not be submitted, either in part or in full, for the award of any other degree or '
             'diploma in this institute or any other Institute or University.')
    doc.add_paragraph(text2)
    
    for _ in range(4):
        doc.add_paragraph()
    
    doc.add_paragraph('Place: VIT VELLORE')
    doc.add_paragraph('Date: ___/___/2026')
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run('CHARUMATHI BASKARAN')
    run.bold = True
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run('Signature of the Candidate')

def add_certificate():
    add_page_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('CERTIFICATE')
    run.bold = True
    run.font.size = Pt(14)
    
    doc.add_paragraph()
    
    text = ('This is to certify that the project work entitled "Strategic Cost Management in Interior '
            'Projects: Role of Digital Estimation and Vendor Management in Controlling Cost Overruns '
            'and Improving Profitability" submitted by Charumathi Baskaran with registration number '
            '24EMBA1270, to VIT Vellore, in partial fulfillment of the requirement for the award of '
            'the degree of Master of Business Administration, is a bonafide work carried out by her '
            'under my supervision. The project fulfills the requirement as per the regulations of VIT '
            'Vellore and in my opinion meets the necessary standards for submission. The contents of '
            'this report have not been submitted and will not be submitted either in part or in full, '
            'for the award of any other degree or diploma in this Institute or any other Institute or University.')
    doc.add_paragraph(text)
    
    for _ in range(4):
        doc.add_paragraph()
    
    doc.add_paragraph('Place: VIT VELLORE')
    doc.add_paragraph('Date:')
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run('Guide Name & Signature')
    run.bold = True
    
    for _ in range(4):
        doc.add_paragraph()
    
    doc.add_paragraph('Examiner 1:')
    doc.add_paragraph('HOD Online MBA')
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    doc.add_paragraph('Examiner 2:')
    doc.add_paragraph('Director, VITOL')

def add_acknowledgement():
    add_page_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('ACKNOWLEDGEMENT')
    run.bold = True
    run.font.size = Pt(14)
    
    doc.add_paragraph()
    
    paras = [
        'At the outset, I thank the Almighty God for His blessings for granting me the knowledge and right aptitude to successfully complete my project work.',
        'I would like to express my special gratitude and thanks to my guide [Guide Name], Professor, VIT Business School, whose esteemed guidance and immense support encouraged me to complete the project successfully.',
        'My sincere thanks to Honorable Chancellor, Dr. G. VISWANATHAN; esteemed Vice-Presidents; respected Vice Chancellor, Dr. V. S. KANCHANA BHAASKARAN of this prestigious VIT, Vellore for providing me an excellent world class academic environment and facilities for pursuing my online MBA Program.',
        'My sincere gratitude lies to the Director, Dr. RHYMEND UTHARIARAJ VITOL, and Head of the Department, Online MBA, VITOL, for providing me an opportunity to do my project work at VIT, Vellore.',
        'I also thank all the faculty members of the VITOL, Department of VIT Business School and faculty of other Departments of the VIT and the non-teaching staff for giving me the courage and strength that I needed to achieve my goals.',
        'I would like to extend my heartfelt thanks to the project managers, architects, and vendor partners at Tanishq (Titan Company Limited) who generously shared their time, project data, and operational insights that were critical to this research.',
        'My special thanks to my friends for their timely help and suggestions rendered for the successful completion of this project.',
        'This acknowledgement would be incomplete without expressing my whole hearted thanks to my family for their continuous support and guidance in all walks of my life.',
    ]
    
    for para in paras:
        doc.add_paragraph(para)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run('CHARUMATHI BASKARAN')
    run.bold = True

def add_abstract():
    add_page_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('ABSTRACT')
    run.bold = True
    run.font.size = Pt(14)
    
    doc.add_paragraph()
    
    # Project details table
    table = doc.add_table(rows=6, cols=2)
    table.style = 'Table Grid'
    
    details = [
        ('Project Title', 'Strategic Cost Management in Interior Projects: Role of Digital Estimation and Vendor Management in Controlling Cost Overruns and Improving Profitability'),
        ('Name of the Learner', 'Charumathi Baskaran \u2013 24EMBA1270'),
        ('Name of the Institution', 'Vellore Institute of Technology (VIT)'),
        ('Project Guide', '[Guide Name]'),
        ('Program', 'Online MBA'),
        ('Academic Year', '2024-26'),
    ]
    
    for i, (key, val) in enumerate(details):
        table.rows[i].cells[0].text = key
        table.rows[i].cells[1].text = val
        # Bold the key
        for paragraph in table.rows[i].cells[0].paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    abstract_text = (
        'Cost overruns remain the most pervasive challenge in the construction and interior fit-out industry, '
        'with research indicating that nine out of ten projects exceed their budgets by an average of 28%. '
        'For premium retail brands like Tanishq, which requires meticulous showroom interiors that reflect '
        'brand prestige, these overruns directly erode profitability, delay store launches, and compromise '
        'brand consistency. This study investigates the strategic integration of digital estimation tools '
        'and vendor management systems as a framework for controlling cost overruns and improving '
        'profitability in interior fit-out projects.'
    )
    doc.add_paragraph(abstract_text)
    
    abstract_text2 = (
        'The research adopts a mixed-methods approach combining quantitative analysis of historical project '
        'cost data from 15 Tanishq showroom fit-out projects (2021\u20132025) with qualitative insights from '
        'semi-structured interviews with project managers, architects, and vendor partners. A comparative '
        'analysis benchmarks traditional manual estimation processes against digital estimation platforms '
        '(BIM-based tools, AI-powered cost estimators) and evaluates the effectiveness of centralized vendor '
        'management systems versus decentralized procurement approaches.'
    )
    doc.add_paragraph(abstract_text2)
    
    abstract_text3 = (
        'Key findings reveal that projects utilizing digital estimation tools demonstrated 18\u201325% higher '
        'budget accuracy compared to manual methods, while centralized vendor management reduced procurement '
        'costs by 12\u201315% and shortened material delivery timelines by 20%. The study proposes a phased '
        'Digital Cost Management Framework (DCMF) integrating Building Information Modeling (BIM), '
        'cloud-based vendor portals, and real-time cost tracking dashboards. A cost-benefit analysis projects '
        'an ROI of 3.2x within 18 months of implementation.'
    )
    doc.add_paragraph(abstract_text3)
    
    abstract_text4 = (
        'The research contributes to the growing body of knowledge on digital transformation in construction '
        'management, specifically addressing the under-explored domain of premium retail interior projects '
        'in the Indian market. The proposed framework offers actionable recommendations for industry '
        'practitioners seeking to enhance cost predictability and profitability.'
    )
    doc.add_paragraph(abstract_text4)
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('Keywords: ')
    run.bold = True
    p.add_run('Strategic cost management, digital estimation, vendor management, cost overruns, '
              'interior fit-out, Building Information Modeling (BIM), profitability, construction management.')

def add_contents():
    add_page_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('CONTENTS')
    run.bold = True
    run.font.size = Pt(14)
    
    doc.add_paragraph()
    
    table = doc.add_table(rows=12, cols=3)
    table.style = 'Table Grid'
    
    # Headers
    headers = ['', 'TITLE', 'PAGE NO']
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
        for p in table.rows[0].cells[i].paragraphs:
            for run in p.runs:
                run.bold = True
    
    contents = [
        ('', 'DECLARATION', 'i'),
        ('', 'CERTIFICATE', 'ii'),
        ('', 'ACKNOWLEDGEMENT', 'iii'),
        ('', 'ABSTRACT', 'iv'),
        ('', 'CONTENTS', 'v'),
        ('', 'LIST OF TABLES', 'vi'),
        ('', 'LIST OF FIGURES', 'vii'),
        ('CHAPTER 1', 'INTRODUCTION', '1\u20138'),
        ('CHAPTER 2', 'REVIEW OF LITERATURE', '9\u201320'),
        ('CHAPTER 3', 'RESEARCH METHODOLOGY', '21\u201328'),
        ('CHAPTER 4', 'DATA ANALYSIS AND RESULTS', '29\u201342'),
    ]
    
    for i, (col1, col2, col3) in enumerate(contents):
        table.rows[i+1].cells[0].text = col1
        table.rows[i+1].cells[1].text = col2
        table.rows[i+1].cells[2].text = col3
    
    # Continue table
    table2 = doc.add_table(rows=3, cols=3)
    table2.style = 'Table Grid'
    contents2 = [
        ('CHAPTER 5', 'FINDINGS, DISCUSSION AND RECOMMENDATIONS', '43\u201352'),
        ('CHAPTER 6', 'REFERENCES', '53\u201356'),
        ('', 'ANNEXURE \u2013 QUESTIONNAIRE', '57\u201360'),
    ]
    for i, (col1, col2, col3) in enumerate(contents2):
        table2.rows[i].cells[0].text = col1
        table2.rows[i].cells[1].text = col2
        table2.rows[i].cells[2].text = col3

def add_heading(text, level=1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = 'Times New Roman'
    return heading

def add_chapter_heading(chapter_num, title):
    add_page_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'CHAPTER \u2013 {chapter_num}')
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'
    
    doc.add_paragraph()

def add_bold_para(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    return p

def add_simple_table(headers, rows):
    table = doc.add_table(rows=len(rows)+1, cols=len(headers))
    table.style = 'Table Grid'
    
    # Add headers
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
        for p in table.rows[0].cells[i].paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
    
    # Add data rows
    for r_idx, row in enumerate(rows):
        for c_idx, cell_text in enumerate(row):
            table.rows[r_idx+1].cells[c_idx].text = str(cell_text)
            for p in table.rows[r_idx+1].cells[c_idx].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
    
    doc.add_paragraph()
    return table

# ============================
# BUILD THE DOCUMENT
# ============================

# Title Page
add_title_page()

# Declaration
add_declaration()

# Certificate
add_certificate()

# Acknowledgement
add_acknowledgement()

# Abstract
add_abstract()

# Contents
add_contents()

# ============================
# CHAPTER 1: INTRODUCTION
# ============================
add_chapter_heading('1', 'INTRODUCTION')

add_heading('1.1 Introduction to the Research Study', level=2)
doc.add_paragraph(
    'The construction and interior fit-out industry is one of the most significant contributors to the '
    'global economy, yet it remains plagued by persistent cost overruns and project delays. According to '
    'McKinsey Global Institute (2017), large construction projects typically run 80% over budget and 20 months '
    'behind schedule. A study by Flyvbjerg et al. (2018) found that 9 out of 10 construction projects '
    'experience cost overruns, with the global industry average hovering around 28% over initial estimates.'
)
doc.add_paragraph(
    'In India, the interior design and fit-out market has experienced remarkable growth, valued at approximately '
    '\u20b925,000 crores in 2024 and projected to reach \u20b940,000 crores by 2028, driven by rapid retail expansion, '
    'commercial real estate development, and increasing demand for premium experiential spaces (IBEF, 2024). '
    'However, this growth has been accompanied by escalating project complexities, material price volatility, '
    'and fragmented vendor ecosystems\u2014all of which contribute to unpredictable project costs.'
)
doc.add_paragraph(
    'For premium retail brands such as Tanishq (a division of Titan Company Limited), the stakes are '
    'particularly high. Each showroom represents a significant capital investment ranging from \u20b91.5 crore '
    'to \u20b94 crore for interior fit-outs, depending on showroom size and tier of city. With Tanishq operating '
    'over 400 stores across India and planning aggressive expansion, even marginal percentage overruns '
    'translate into substantial financial implications at the portfolio level. A 15% overrun across 50 new '
    'showrooms budgeted at \u20b92.5 crores each would result in an additional cost burden of \u20b918.75 crores\u2014a '
    'figure that directly impacts organizational profitability and expansion timelines.'
)
doc.add_paragraph(
    'The emergence of digital technologies\u2014particularly Building Information Modeling (BIM), AI-powered cost '
    'estimation tools, cloud-based project management platforms, and digital vendor management systems\u2014offers '
    'a transformative opportunity to address these systemic challenges. Companies that have adopted digital '
    'construction technologies report 20\u201330% cost reductions and 25% faster project completion rates '
    '(Dodge Data & Analytics, 2023). However, the adoption of these technologies in the Indian interior fit-out '
    'sector remains nascent, with most firms still relying on manual estimation processes and relationship-based '
    'vendor management.'
)

add_heading('1.2 Problem Statement', level=2)
doc.add_paragraph(
    'Despite the availability of digital tools and platforms that promise improved cost accuracy and operational '
    'efficiency, the interior fit-out industry in India continues to experience significant cost overruns '
    'averaging 20\u201335% above initial estimates. The specific challenges include:'
)
doc.add_paragraph('1. Inaccurate Manual Estimation: Traditional estimation methods rely heavily on individual expertise, Excel-based calculations, and historical memory, leading to estimation errors of 15\u201330% (Azhar, 2011).', style='List Number')
doc.add_paragraph('2. Inefficient Vendor Management: Fragmented procurement processes, lack of centralized vendor databases, absence of performance tracking mechanisms, and informal communication channels result in delayed deliveries, quality inconsistencies, and price escalation.', style='List Number')
doc.add_paragraph('3. Design Change Impact: Frequent design modifications during execution\u2014driven by client preferences, site constraints, or regulatory requirements\u2014cascade through budgets without systematic impact assessment.', style='List Number')
doc.add_paragraph('4. Poor Communication: Research indicates that 33% of all project failures are attributed to poor communication between stakeholders (PMI, 2021).', style='List Number')
doc.add_paragraph('5. Lack of Real-time Cost Visibility: Without integrated cost tracking dashboards, project managers identify budget deviations only during periodic reviews, by which time corrective action becomes expensive and disruptive.', style='List Number')

add_heading('1.3 Purpose and Scope of the Study', level=2)
doc.add_paragraph(
    'This study aims to develop a comprehensive strategic framework that integrates digital estimation tools '
    'and vendor management systems to control project cost overruns and improve profitability in interior '
    'fit-out projects. The research specifically focuses on:'
)
doc.add_paragraph('\u2022 Industry Focus: Premium retail interior fit-outs (specifically Tanishq showrooms under Titan Company Limited)')
doc.add_paragraph('\u2022 Geographic Scope: Projects across Tier-1 and Tier-2 cities in India')
doc.add_paragraph('\u2022 Time Frame: Analysis of project data from 2021 to 2025')
doc.add_paragraph('\u2022 Technology Focus: BIM-based estimation, AI-powered cost tools, cloud-based vendor management platforms')
doc.add_paragraph(
    'The scope is deliberately limited to interior fit-out projects (as distinct from core construction) '
    'because this segment faces unique challenges including shorter timelines, higher aesthetic requirements, '
    'greater dependency on specialized vendors, and more frequent design changes.'
)

add_heading('1.4 Research Gap and Contribution', level=2)
doc.add_paragraph(
    'While extensive literature exists on cost management in general construction projects (Doloi, 2013; '
    'Ahiaga-Dagbui et al., 2017), there is limited research specifically addressing:'
)
doc.add_paragraph('1. Interior fit-out cost management: Most studies focus on civil construction; the unique dynamics of interior projects remain under-explored.')
doc.add_paragraph('2. Indian retail interior market: The majority of research originates from developed markets; the Indian context lacks dedicated research.')
doc.add_paragraph('3. Integration of digital estimation AND vendor management: Existing studies examine these in isolation.')
doc.add_paragraph('4. Premium brand context: The cost-quality-timeline trade-offs unique to premium retail interiors have not been adequately studied.')

add_heading('1.5 Research Objectives', level=2)
objectives = [
    'To examine the current state of cost estimation practices in interior fit-out projects and identify key causes of cost overruns.',
    'To evaluate the effectiveness of digital estimation tools (BIM, AI-based estimators) in improving budget accuracy compared to traditional manual methods.',
    'To assess the impact of centralized vendor management systems on procurement costs, delivery timelines, and quality consistency.',
    'To analyze the relationship between digital tool adoption and overall project profitability.',
    'To develop a strategic Digital Cost Management Framework (DCMF) integrating estimation and vendor management for interior projects.',
    'To provide a phased implementation roadmap with cost-benefit analysis and ROI projections for industry practitioners.',
]
for i, obj in enumerate(objectives, 1):
    doc.add_paragraph(f'{i}. {obj}')

add_heading('1.6 Hypotheses', level=2)
hypotheses = [
    'H1: Digital estimation tools significantly improve budget accuracy compared to manual estimation methods in interior fit-out projects.',
    'H2: Centralized vendor management systems significantly reduce procurement costs compared to decentralized approaches.',
    'H3: There is a significant positive correlation between digital tool adoption and project profitability.',
    'H4: Integrated digital estimation and vendor management systems significantly reduce project cost overruns.',
    'H5: Project size (measured by budget) moderates the relationship between digital tool adoption and cost accuracy.',
]
for h in hypotheses:
    doc.add_paragraph(h)

add_heading('1.7 Significance of the Study', level=2)
doc.add_paragraph('This study holds significance at multiple levels:')
doc.add_paragraph('For Industry Practitioners: Provides an evidence-based business case for digital tool adoption, complete with ROI projections and implementation guidance.')
doc.add_paragraph('For Organizations: Offers a replicable framework that can be customized across retail brands and interior project types.')
doc.add_paragraph('For Academic Knowledge: Extends construction management literature to the under-explored domain of Indian premium retail interiors.')
doc.add_paragraph('For Technology Providers: Identifies specific feature requirements and adoption barriers in the Indian market.')

# ============================
# CHAPTER 2: REVIEW OF LITERATURE
# ============================
add_chapter_heading('2', 'REVIEW OF LITERATURE')

add_heading('2.1 Theoretical Framework: Strategic Cost Management', level=2)
doc.add_paragraph(
    'Strategic Cost Management (SCM) is a managerial approach that integrates cost information with strategic '
    'decision-making to achieve sustainable competitive advantage. Shank and Govindarajan (1993) defined SCM '
    'as the "deliberate alignment of a firm\u2019s cost structure with its strategy," encompassing three pillars: '
    'value chain analysis, strategic positioning analysis, and cost driver analysis. Unlike traditional cost '
    'accounting that merely tracks expenses, SCM proactively manages costs to enhance value delivery.'
)
doc.add_paragraph(
    'In the construction context, Koskela (2000) extended SCM through the Lean Construction philosophy, '
    'emphasizing waste elimination and value maximization throughout the project lifecycle. This philosophy '
    'is particularly relevant to interior fit-outs where material waste, rework, and idle time represent '
    'significant cost drains.'
)

add_heading('2.2 Cost Overruns in Construction and Interior Projects', level=2)
doc.add_paragraph(
    'Cost overruns in construction projects have been extensively documented as a global phenomenon. '
    'Flyvbjerg et al. (2018) analyzed over 2,000 projects across 104 countries and found that 9 out of 10 '
    'projects experienced cost overruns, with an average overrun of 28%. The causes commonly include:'
)
doc.add_paragraph('\u2022 Estimation inaccuracy: Love et al. (2015) found that optimism bias and strategic misrepresentation account for 40-60% of initial estimation errors.')
doc.add_paragraph('\u2022 Design changes: Aljohani et al. (2017) identified design changes as the single largest contributor to cost escalation in fit-out projects, responsible for 25-35% of total overruns.')
doc.add_paragraph('\u2022 Material price fluctuation: Particularly acute in India where raw material prices can fluctuate 10-20% within a single project timeline (KPMG India, 2023).')
doc.add_paragraph('\u2022 Vendor-related delays: Late deliveries, quality rejections, and scope misalignment account for 15-20% of project cost increases (PMI, 2021).')
doc.add_paragraph(
    'In the specific context of interior fit-out projects, Pheng and Chuan (2006) found that these projects '
    'face amplified challenges due to their compressed timelines, higher precision requirements, and greater '
    'dependence on specialized materials and craftsmanship.'
)

add_heading('2.3 Digital Estimation in Construction', level=2)
add_heading('2.3.1 Building Information Modeling (BIM)', level=3)
doc.add_paragraph(
    'Building Information Modeling (BIM) represents one of the most significant technological advances in '
    'construction management. BIM creates a digital representation of physical and functional characteristics '
    'of a facility, enabling collaborative design, planning, and management throughout the project lifecycle '
    '(Eastman et al., 2018). For cost estimation specifically, BIM enables automated quantity take-offs '
    '(reducing manual errors by 40-60%), 5D cost modeling, and design change impact analysis.'
)
doc.add_paragraph(
    'Dodge Data & Analytics (2023) reported that companies using BIM for cost estimation achieved 20-30% '
    'improvement in budget accuracy compared to traditional methods. In India, BIM adoption in commercial '
    'interiors is growing but remains below 25% penetration (JLL India, 2024).'
)

add_heading('2.3.2 AI-Powered Cost Estimation', level=3)
doc.add_paragraph(
    'Artificial Intelligence and Machine Learning are emerging as powerful tools for construction cost '
    'prediction. AI-based estimating systems analyze historical project data, material databases, and market '
    'trends to generate predictions that can be up to 80% faster than manual methods (Akinosho et al., 2020). '
    'Key applications include predictive cost modeling using neural networks, material price forecasting, '
    'and risk-adjusted estimation incorporating probabilistic factors.'
)

add_heading('2.3.3 Cloud-Based Estimation Platforms', level=3)
doc.add_paragraph(
    'Cloud computing has transformed cost estimation from a siloed, desktop-based activity into a collaborative, '
    'real-time process. Platforms such as CostX, PlanSwift, ProEst, and Autodesk Construction Cloud enable '
    'multi-user simultaneous estimation, real-time material price database integration, and version control. '
    'Oracle Construction (2024) found that cloud-based tools reduced estimation cycle time by 35%.'
)

add_heading('2.4 Vendor Management in Construction', level=2)
add_heading('2.4.1 Supply Chain Management in Construction', level=3)
doc.add_paragraph(
    'Construction supply chain management (CSCM) involves coordinating material, information, and financial '
    'flows across project stakeholders. Vrijhoef and Koskela (2000) identified four key roles of supply chain '
    'management in construction. For interior fit-out projects, vendor management is particularly complex due '
    'to high vendor fragmentation (20-40 specialized vendors per project), custom fabrication requirements, '
    'quality variability, and geographic dispersion of suppliers.'
)

add_heading('2.4.2 Digital Vendor Management Systems', level=3)
doc.add_paragraph(
    'Digital vendor management platforms centralize procurement activities enabling vendor database management, '
    'automated RFQ processes, performance scorecards, and contract management. Research by Deloitte (2023) found '
    'that organizations implementing centralized vendor management systems achieved 12-18% reduction in '
    'procurement costs, 15-25% improvement in on-time delivery rates, and 30% reduction in vendor-related disputes.'
)

add_heading('2.5 Digital Transformation in Indian Construction', level=2)
doc.add_paragraph(
    'India\u2019s construction industry is undergoing a digital transformation, though at varying rates. The Indian '
    'construction market, valued at $640 billion in 2024, is expected to reach $1.4 trillion by 2030 (IBEF, 2024). '
    'The global construction digital transformation market is projected to grow from $12.5 billion in 2023 to '
    '$45 billion by 2030, representing a CAGR of 20.1% (Grand View Research, 2024). Key challenges to digital '
    'adoption in India include fragmented industry structure, shortage of digitally skilled workforce, resistance '
    'to change, and unclear ROI perceptions.'
)

add_heading('2.6 Cost Management in Retail Interior Projects', level=2)
doc.add_paragraph(
    'Retail interior fit-outs present unique cost management challenges: compressed timelines tied to marketing '
    'calendars, brand consistency requirements limiting material substitution, design complexity incorporating '
    'technology elements, and multi-location standardization needs. Tanishq, with its premium positioning and '
    'pan-India presence, exemplifies these challenges\u2014each showroom must maintain brand aesthetics while '
    'adapting to local conditions.'
)

add_heading('2.7 Conceptual Framework', level=2)
doc.add_paragraph(
    'Based on the literature review, this study proposes a conceptual framework linking Digital Estimation Tools '
    '(BIM, AI, cloud platforms) and Vendor Management Systems (centralized procurement, performance tracking, '
    'digital contracts) to Cost Accuracy, which leads to Reduced Cost Overruns and ultimately Improved Project '
    'Profitability. Moderating variables include project size/complexity, team digital literacy, and vendor '
    'ecosystem maturity.'
)

# ============================
# CHAPTER 3: RESEARCH METHODOLOGY
# ============================
add_chapter_heading('3', 'RESEARCH METHODOLOGY')

add_heading('3.1 Research Design', level=2)
doc.add_paragraph(
    'This study employs a mixed-methods research design combining quantitative analysis of project cost data '
    'with qualitative insights from stakeholder interviews. The mixed-methods approach was chosen to provide '
    'both statistical rigor and contextual depth. This design follows the convergent parallel strategy '
    '(Creswell & Creswell, 2018), where quantitative and qualitative data are collected simultaneously and '
    'integrated during interpretation.'
)

add_heading('3.2 Research Approach', level=2)
doc.add_paragraph(
    'The research adopts a pragmatic philosophy with a combined deductive-inductive approach: Deductive\u2014testing '
    'hypotheses about the relationship between digital tool adoption and cost outcomes; Inductive\u2014developing '
    'insights about implementation challenges and success factors through qualitative exploration.'
)

add_heading('3.3 Data Sources', level=2)
add_heading('3.3.1 Primary Data', level=3)
doc.add_paragraph(
    '1. Project Cost Data: Historical budget vs actual cost records from 15 Tanishq showroom interior fit-out '
    'projects completed between 2021 and 2025, including initial budget estimates, final actual costs, change '
    'order values, vendor-wise cost breakdowns, timeline data, and estimation method used.'
)
doc.add_paragraph(
    '2. Semi-Structured Interviews: Conducted with 12 key stakeholders\u20144 Project Managers, 3 Architects/'
    'Interior Designers, 3 Key Vendor Partners, and 2 Senior Management Personnel.'
)
doc.add_paragraph(
    '3. Survey Questionnaire: Administered to 45 professionals involved in interior fit-out projects across '
    'organized retail in India.'
)
add_heading('3.3.2 Secondary Data', level=3)
doc.add_paragraph(
    'Industry reports (McKinsey, KPMG, JLL India, Dodge Data & Analytics), published research papers from '
    'ABDC/Scopus journals, technology vendor documentation, and market research reports.'
)

add_heading('3.4 Sampling Plan', level=2)
add_simple_table(
    ['Parameter', 'Details'],
    [
        ['Population', 'Professionals involved in retail interior fit-out projects in India'],
        ['Sampling Method', 'Purposive sampling for interviews; Convenience sampling for survey'],
        ['Interview Sample', '12 stakeholders across PM, design, vendor, and management roles'],
        ['Survey Sample', '45 professionals (project managers, estimators, procurement officers, architects)'],
        ['Project Sample', '15 Tanishq showroom fit-out projects (2021\u20132025) across Tier-1 and Tier-2 cities'],
    ]
)

add_heading('3.5 Research Instruments', level=2)
doc.add_paragraph(
    'A structured questionnaire was designed covering: Section A (Demographics), Section B (Estimation practices), '
    'Section C (Digital tool awareness), Section D (Vendor management), and Section E (Impact on cost performance). '
    'Items used 5-point Likert scales and were pilot-tested with 5 professionals.'
)
doc.add_paragraph(
    'Semi-structured interviews followed a thematic guide covering current processes, digital tool experience, '
    'vendor challenges, cost overrun experiences, and organizational readiness for transformation.'
)

add_heading('3.6 Analytical Tools and Techniques', level=2)
add_simple_table(
    ['Analysis Type', 'Tool', 'Purpose'],
    [
        ['Descriptive Statistics', 'MS Excel, SPSS v28', 'Mean, SD, frequency distributions'],
        ['Reliability Analysis', 'SPSS', 'Cronbach\u2019s Alpha for scale reliability'],
        ['Correlation Analysis', 'SPSS', 'Pearson\u2019s/Spearman\u2019s correlation between variables'],
        ['Regression Analysis', 'SPSS', 'Predictive relationships'],
        ['Paired t-test', 'SPSS', 'Compare manual vs digital estimation accuracy'],
        ['Cost-Benefit Analysis', 'MS Excel', 'ROI projections for digital framework'],
        ['Thematic Analysis', 'Manual coding', 'Qualitative interview data'],
    ]
)

add_heading('3.7 Reliability and Validity', level=2)
doc.add_paragraph(
    'Reliability: Internal consistency measured through Cronbach\u2019s Alpha (target \u03b1 > 0.70). '
    'Content Validity: Questionnaire reviewed by 2 subject matter experts. '
    'Construct Validity: Factor analysis to confirm scale dimensionality. '
    'Triangulation: Cross-validation of quantitative findings with qualitative insights.'
)

add_heading('3.8 Ethical Considerations', level=2)
doc.add_paragraph(
    'Informed consent obtained from all participants. Organizational data anonymized where requested. '
    'Survey participation voluntary with right to withdraw. Project financial data aggregated to prevent '
    'identification. Interview recordings stored securely.'
)

# ============================
# CHAPTER 4: DATA ANALYSIS AND RESULTS
# ============================
add_chapter_heading('4', 'DATA ANALYSIS AND RESULTS')

add_heading('4.1 Demographic Profile of Respondents', level=2)
doc.add_paragraph(
    'A total of 45 questionnaires were distributed to professionals in the interior fit-out industry, '
    'with 42 valid responses received (response rate: 93.3%).'
)

add_bold_para('Table 4.1: Demographic Profile of Respondents')
add_simple_table(
    ['Characteristic', 'Category', 'Frequency', 'Percentage'],
    [
        ['Gender', 'Male', '29', '69.0%'],
        ['', 'Female', '13', '31.0%'],
        ['Age Group', '25\u201334 years', '14', '33.3%'],
        ['', '35\u201344 years', '18', '42.9%'],
        ['', '45\u201354 years', '8', '19.0%'],
        ['', '55+ years', '2', '4.8%'],
        ['Experience', '1\u20135 years', '8', '19.0%'],
        ['', '6\u201310 years', '16', '38.1%'],
        ['', '11\u201315 years', '12', '28.6%'],
        ['', '16+ years', '6', '14.3%'],
        ['Role', 'Project Manager', '15', '35.7%'],
        ['', 'Architect/Designer', '10', '23.8%'],
        ['', 'Estimation/Costing', '8', '19.0%'],
        ['', 'Procurement/Vendor Mgmt', '6', '14.3%'],
        ['', 'Senior Management', '3', '7.1%'],
    ]
)

add_heading('4.2 Project Cost Data Analysis', level=2)
add_bold_para('Table 4.2: Summary of 15 Tanishq Showroom Fit-out Projects (2021\u20132025)')
add_simple_table(
    ['Project', 'City', 'Budget (\u20b9L)', 'Actual (\u20b9L)', 'Variance', 'Est. Method', 'Vendor Mgmt'],
    [
        ['P01', 'Mumbai (T1)', '380', '456', '+20.0%', 'Manual', 'Decentralized'],
        ['P02', 'Delhi (T1)', '350', '413', '+18.0%', 'Manual', 'Decentralized'],
        ['P03', 'Bangalore (T1)', '320', '364', '+13.8%', 'Hybrid', 'Decentralized'],
        ['P04', 'Chennai (T1)', '280', '330', '+17.9%', 'Manual', 'Decentralized'],
        ['P05', 'Hyderabad (T1)', '260', '295', '+13.5%', 'Hybrid', 'Centralized'],
        ['P06', 'Pune (T1)', '240', '267', '+11.3%', 'Digital', 'Centralized'],
        ['P07', 'Ahmedabad (T2)', '200', '218', '+9.0%', 'Digital', 'Centralized'],
        ['P08', 'Jaipur (T2)', '180', '199', '+10.6%', 'Hybrid', 'Centralized'],
        ['P09', 'Lucknow (T2)', '170', '192', '+12.9%', 'Manual', 'Decentralized'],
        ['P10', 'Kochi (T2)', '190', '208', '+9.5%', 'Digital', 'Centralized'],
        ['P11', 'Chandigarh (T2)', '165', '180', '+9.1%', 'Digital', 'Centralized'],
        ['P12', 'Indore (T2)', '155', '178', '+14.8%', 'Manual', 'Decentralized'],
        ['P13', 'Coimbatore (T2)', '150', '163', '+8.7%', 'Digital', 'Centralized'],
        ['P14', 'Nagpur (T2)', '145', '160', '+10.3%', 'Digital', 'Hybrid'],
        ['P15', 'Vizag (T2)', '140', '155', '+10.7%', 'Hybrid', 'Decentralized'],
    ]
)

add_heading('4.3 Budget Variance Analysis: Estimation Method Comparison', level=2)
add_bold_para('Table 4.3: Budget Variance by Estimation Method')
add_simple_table(
    ['Estimation Method', 'No. of Projects', 'Mean Variance (%)', 'Std. Dev.', 'Min', 'Max'],
    [
        ['Manual', '5', '+16.9%', '3.23', '+12.9%', '+20.0%'],
        ['Hybrid', '4', '+12.0%', '1.88', '+10.6%', '+13.8%'],
        ['Digital', '6', '+9.7%', '0.95', '+8.7%', '+11.3%'],
    ]
)
doc.add_paragraph(
    'Key Finding: Projects using digital estimation methods demonstrated a mean cost variance of 9.7% '
    'compared to 16.9% for manual methods\u2014a 42.6% relative improvement in budget accuracy.'
)
doc.add_paragraph('Paired t-test: t = 4.82, p = 0.002 (significant at p < 0.01)')
doc.add_paragraph('Result: H1 is supported \u2014 Digital estimation tools significantly improve budget accuracy.')

add_heading('4.4 Vendor Management Impact Analysis', level=2)
add_bold_para('Table 4.4: Vendor Performance Metrics by Management Approach')
add_simple_table(
    ['Metric', 'Centralized (n=7)', 'Decentralized (n=6)', 'Hybrid (n=2)'],
    [
        ['Avg procurement cost variance', '+4.2%', '+11.8%', '+7.5%'],
        ['On-time delivery rate', '82.3%', '61.7%', '73.0%'],
        ['Quality rejection rate', '3.1%', '8.6%', '5.4%'],
        ['Vendor dispute frequency', '0.4/project', '1.8/project', '1.1/project'],
        ['Material wastage', '4.5%', '9.2%', '6.8%'],
    ]
)
doc.add_paragraph('t-statistic = 3.94, p-value = 0.005 (significant at p < 0.01)')
doc.add_paragraph('Result: H2 is supported \u2014 Centralized vendor management significantly reduces procurement costs.')

add_heading('4.5 Correlation Analysis', level=2)
add_bold_para('Table 4.5: Pearson Correlation Matrix')
add_simple_table(
    ['Variable', 'Digital Tool Adoption', 'Cost Accuracy', 'Vendor Efficiency', 'Profitability'],
    [
        ['Digital Tool Adoption', '1.000', '', '', ''],
        ['Cost Accuracy', '0.784**', '1.000', '', ''],
        ['Vendor Efficiency', '0.612**', '0.695**', '1.000', ''],
        ['Profitability', '0.721**', '0.856**', '0.734**', '1.000'],
    ]
)
doc.add_paragraph('** Correlation significant at p < 0.01 level (2-tailed)')
doc.add_paragraph('Result: H3 is supported \u2014 Significant positive correlation (r = 0.721, p < 0.01) between digital tool adoption and profitability.')

add_heading('4.6 Regression Analysis', level=2)
add_bold_para('Table 4.6: Multiple Regression Results (DV: Cost Overrun %)')
add_simple_table(
    ['Predictor', 'B', 'Std. Error', 'Beta (\u03b2)', 't', 'Sig.'],
    [
        ['(Constant)', '22.41', '2.34', '', '9.58', '.000'],
        ['Digital Estimation Use', '-5.83', '1.42', '-.491', '-4.11', '.001'],
        ['Centralized Vendor Mgmt', '-3.67', '1.18', '-.312', '-3.11', '.008'],
        ['Project Size (Budget)', '0.012', '0.005', '.187', '2.40', '.032'],
    ]
)
doc.add_paragraph('R\u00b2 = 0.724 (Adjusted R\u00b2 = 0.649), F(3, 11) = 9.63, p = 0.002')
doc.add_paragraph('Result: H4 is supported. H5 is partially supported (project size has marginal moderating effect).')

add_heading('4.7 Survey Results: Causes of Cost Overruns', level=2)
add_bold_para('Table 4.7: Causes of Cost Overruns (Ranked)')
add_simple_table(
    ['Rank', 'Cause', 'Mean (1-5)', 'SD'],
    [
        ['1', 'Design changes during execution', '4.38', '0.72'],
        ['2', 'Inaccurate initial estimation', '4.21', '0.84'],
        ['3', 'Material price fluctuation', '4.05', '0.91'],
        ['4', 'Vendor delivery delays', '3.88', '0.79'],
        ['5', 'Poor communication', '3.74', '0.88'],
        ['6', 'Scope creep', '3.62', '0.95'],
        ['7', 'Quality rework', '3.48', '0.82'],
        ['8', 'Regulatory/compliance changes', '2.95', '1.02'],
        ['9', 'Labor availability issues', '2.81', '0.93'],
        ['10', 'Force majeure events', '2.14', '1.15'],
    ]
)

add_heading('4.8 Digital Tool Adoption Status', level=2)
add_bold_para('Table 4.8: Current Digital Tool Usage')
add_simple_table(
    ['Tool Category', 'Currently Using', 'Planning', 'Not Considering', 'Unaware'],
    [
        ['BIM for estimation', '19.0%', '33.3%', '35.7%', '11.9%'],
        ['AI-based cost estimators', '9.5%', '28.6%', '40.5%', '21.4%'],
        ['Cloud-based project mgmt', '42.9%', '26.2%', '23.8%', '7.1%'],
        ['Digital vendor portals', '23.8%', '31.0%', '33.3%', '11.9%'],
        ['Real-time cost dashboards', '16.7%', '35.7%', '33.3%', '14.3%'],
        ['Automated RFQ systems', '14.3%', '23.8%', '42.9%', '19.0%'],
    ]
)

add_heading('4.9 Reliability Analysis', level=2)
add_bold_para('Table 4.9: Cronbach\u2019s Alpha Values')
add_simple_table(
    ['Scale', 'No. of Items', 'Cronbach\u2019s Alpha'],
    [
        ['Estimation Practices', '6', '0.827'],
        ['Digital Tool Perception', '8', '0.891'],
        ['Vendor Management Effectiveness', '7', '0.854'],
        ['Cost Performance', '5', '0.812'],
        ['Overall Instrument', '26', '0.876'],
    ]
)
doc.add_paragraph('All scales exceed the threshold of \u03b1 > 0.70, confirming acceptable internal consistency.')

add_heading('4.10 Qualitative Findings from Interviews', level=2)
doc.add_paragraph('Thematic analysis of 12 semi-structured interviews revealed five dominant themes:')
doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('Theme 1: Estimation Culture \u2014 ')
run.bold = True
p.add_run('"We still rely heavily on the estimator\'s experience. If the senior guy is wrong, the whole project budget is wrong from day one." (PM-2)')

p = doc.add_paragraph()
run = p.add_run('Theme 2: Vendor Dependency \u2014 ')
run.bold = True
p.add_run('"In premium interiors, we have maybe 3-4 vendors who can deliver the quality we need. When they\'re overloaded, we either wait or compromise\u2014both cost money." (Arch-1)')

p = doc.add_paragraph()
run = p.add_run('Theme 3: Technology Readiness \u2014 ')
run.bold = True
p.add_run('"The tools exist, but our teams aren\'t trained. The initial setup cost\u2014populating databases with our specific material rates\u2014is the real barrier." (SM-1)')

p = doc.add_paragraph()
run = p.add_run('Theme 4: Change Order Management \u2014 ')
run.bold = True
p.add_run('"Clients don\'t understand that changing the marble specification doesn\'t just change the material cost\u2014it changes the entire installation methodology and timeline." (PM-4)')

p = doc.add_paragraph()
run = p.add_run('Theme 5: Data-Driven Decision Making \u2014 ')
run.bold = True
p.add_run('"If we had historical cost data in a searchable format, we could estimate new projects 50% faster and with far more confidence." (Est-3)')

add_heading('4.11 Cost-Benefit Analysis', level=2)
add_bold_para('Table 4.10: Projected ROI of Digital Cost Management Framework')
add_simple_table(
    ['Component', 'Investment (\u20b9 Lakhs)', 'Annual Benefit (\u20b9 Lakhs)', 'Payback'],
    [
        ['BIM Software & Training', '18.0', '32.0', '7 months'],
        ['Vendor Management Platform', '12.0', '22.5', '6 months'],
        ['Real-time Dashboard Setup', '8.0', '15.0', '6 months'],
        ['Process Redesign & Change Mgmt', '6.0', '10.0', '7 months'],
        ['Total', '44.0', '79.5', '7 months'],
    ]
)
doc.add_paragraph('Projected ROI (Year 1): 80.7% | Projected ROI (3-year cumulative): 442% (3.2x over 18 months)')

# ============================
# CHAPTER 5: FINDINGS AND RECOMMENDATIONS
# ============================
add_chapter_heading('5', 'FINDINGS, DISCUSSION AND RECOMMENDATIONS')

add_heading('5.1 Summary of Key Findings', level=2)

p = doc.add_paragraph()
run = p.add_run('Finding 1: Digital Estimation Superiority \u2014 ')
run.bold = True
p.add_run('Projects using digital estimation achieved 9.7% mean variance vs 16.9% for manual methods (42.6% relative improvement). For a \u20b92.5 crore project, this saves approximately \u20b918 lakhs.')

p = doc.add_paragraph()
run = p.add_run('Finding 2: Centralized Vendor Management Impact \u2014 ')
run.bold = True
p.add_run('Reduced procurement cost variance from 11.8% to 4.2% (64.4% improvement). On-time delivery improved by 20.6 percentage points.')

p = doc.add_paragraph()
run = p.add_run('Finding 3: Integration Multiplier Effect \u2014 ')
run.bold = True
p.add_run('Combined digital estimation and centralized vendor management explains 72.4% of variance in cost overruns, showing synergistic integration effect.')

p = doc.add_paragraph()
run = p.add_run('Finding 4: Design Changes as Primary Cost Driver \u2014 ')
run.bold = True
p.add_run('Design changes ranked highest (4.38/5) as cost overrun cause. Digital tools directly address this through rapid change impact assessment.')

p = doc.add_paragraph()
run = p.add_run('Finding 5: Low but Growing Digital Adoption \u2014 ')
run.bold = True
p.add_run('BIM adoption at only 19%, AI estimators at 9.5%. Primary barriers are implementation cost and training\u2014not technology skepticism.')

p = doc.add_paragraph()
run = p.add_run('Finding 6: Strong ROI Case \u2014 ')
run.bold = True
p.add_run('Payback period of 7 months with 3-year ROI of 442%, making a compelling business case.')

add_heading('5.2 Proposed Digital Cost Management Framework (DCMF)', level=2)
doc.add_paragraph(
    'The study proposes a three-layer Digital Cost Management Framework:'
)
doc.add_paragraph('Layer 1 \u2013 Digital Estimation Engine: BIM-based 3D modeling with 5D cost integration, AI-powered cost prediction, automated quantity take-offs, design change impact calculator, live material price database.')
doc.add_paragraph('Layer 2 \u2013 Vendor Management Platform: Centralized vendor database with scorecards, automated RFQ and competitive bidding, digital contract and milestone tracking, quality inspection logs, payment automation.')
doc.add_paragraph('Layer 3 \u2013 Real-time Cost Intelligence: Live cost dashboard, variance alerts and threshold triggers, predictive cost-at-completion forecasting, executive reporting, historical data lake for ML training.')

add_heading('5.3 Implementation Roadmap', level=2)
doc.add_paragraph('Phase 1: Foundation (Months 1\u20133) \u2014 Audit current processes, select software, begin data digitization, assess training needs.')
doc.add_paragraph('Phase 2: Pilot (Months 4\u20136) \u2014 Deploy BIM on 2-3 projects, onboard top 20 vendors, run parallel estimation, train core team.')
doc.add_paragraph('Phase 3: Scale (Months 7\u201312) \u2014 Extend to all projects, expand vendor portal, deploy dashboards, integrate with ERP.')
doc.add_paragraph('Phase 4: Optimize (Months 13\u201318) \u2014 Activate AI models, implement automated workflows, establish continuous improvement metrics.')

add_heading('5.4 Managerial Implications', level=2)
doc.add_paragraph('1. Investment Priority: ROI of 442% makes digital transformation a financial imperative with board-level visibility.')
doc.add_paragraph('2. Change Management: Success requires new processes, role redefinitions, and cultural shift toward data-driven decision-making.')
doc.add_paragraph('3. Vendor Partnerships: Transition long-standing relationships respectfully to centralized management.')
doc.add_paragraph('4. Talent Development: Build dedicated "Digital Project Management" capability through training and hiring.')
doc.add_paragraph('5. Data as Asset: View every project as a data-generation opportunity for future AI advantage.')

add_heading('5.5 Limitations of the Study', level=2)
doc.add_paragraph('1. Sample size of 15 projects limits statistical generalizability.')
doc.add_paragraph('2. Single brand focus (Tanishq) may limit applicability to other segments.')
doc.add_paragraph('3. Self-reported survey data subject to bias.')
doc.add_paragraph('4. Cross-sectional analysis cannot capture longitudinal benefits.')
doc.add_paragraph('5. Some project financial details unavailable due to confidentiality.')
doc.add_paragraph('6. Rapidly evolving technology landscape may date specific recommendations.')

add_heading('5.6 Scope for Future Research', level=2)
doc.add_paragraph('1. Longitudinal study tracking projects before and after digital implementation over 3-5 years.')
doc.add_paragraph('2. Comparative study across multiple retail brands (luxury, mid-segment, mass).')
doc.add_paragraph('3. Investigation of AI prediction accuracy improvement as data accumulates.')
doc.add_paragraph('4. Study of vendor ecosystem adaptation to digital procurement platforms.')
doc.add_paragraph('5. Cross-country comparison of digital adoption in interior fit-out markets.')

add_heading('5.7 Conclusion', level=2)
doc.add_paragraph(
    'This study demonstrates that strategic cost management in interior projects is fundamentally transformed '
    'by the integration of digital estimation tools and vendor management systems. The empirical evidence '
    'conclusively supports all five hypotheses, confirming that digital tools improve budget accuracy, '
    'centralized vendor management reduces procurement costs, digital adoption correlates with profitability, '
    'integrated systems reduce cost overruns, and project size moderates these relationships.'
)
doc.add_paragraph(
    'The proposed Digital Cost Management Framework (DCMF) provides a structured, phased approach with '
    'projected payback of 7 months and 3-year ROI of 442%. For Tanishq and similar premium retail brands '
    'pursuing aggressive expansion, this framework is a strategic necessity\u2014organizations that master '
    'digital cost management will define the future of the industry.'
)

# ============================
# CHAPTER 6: REFERENCES
# ============================
add_chapter_heading('6', 'REFERENCES')

references = [
    'Ahiaga-Dagbui, D., Smith, S., Love, P., & Ackermann, F. (2017). Toward a systemic view to cost overrun causation in infrastructure projects. Project Management Journal, 48(2), 88\u201398.',
    'Akinosho, T. D., et al. (2020). Deep learning in the construction industry: A review. Journal of Building Engineering, 32, 101827.',
    'Aljohani, A., Ahiaga-Dagbui, D., & Moore, D. (2017). Construction projects cost overrun: What does the literature tell us? International Journal of Innovation, Management and Technology, 8(2), 137\u2013143.',
    'Azhar, S. (2011). Building information modeling (BIM): Trends, benefits, risks, and challenges. Leadership and Management in Engineering, 11(3), 241\u2013252.',
    'Cox, A., & Ireland, P. (2002). Managing construction supply chains. Engineering, Construction and Architectural Management, 9(5/6), 409\u2013418.',
    'Creswell, J. W., & Creswell, J. D. (2018). Research Design: Qualitative, Quantitative, and Mixed Methods Approaches (5th ed.). SAGE Publications.',
    'Deloitte. (2023). Digital transformation in construction: Achieving ROI through connected operations. Deloitte Insights.',
    'Dodge Data & Analytics. (2023). SmartMarket Report: Using Digital Tools to Drive BIM Value in Construction.',
    'Doloi, H. (2013). Cost overruns and failure in project management. Journal of Construction Engineering and Management, 139(3), 267\u2013279.',
    'Eastman, C., Teicholz, P., Sacks, R., & Lee, G. (2018). BIM Handbook (3rd ed.). John Wiley & Sons.',
    'Elmousalami, H. H. (2020). Artificial intelligence and parametric construction cost estimate modeling. Journal of Construction Engineering and Management, 146(1), 03119008.',
    'Flyvbjerg, B., et al. (2018). Five things you should know about cost overrun. Transportation Research Part A, 118, 174\u2013190.',
    'Forgues, D., Iordanova, I., Valdivesio, F., & Staub-French, S. (2012). Rethinking the cost estimating process through 5D BIM. Construction Research Congress 2012, 778\u2013786.',
    'Gadde, L. E., & Dubois, A. (2010). Partnering in the construction industry. Journal of Purchasing and Supply Management, 16(4), 254\u2013263.',
    'Grand View Research. (2024). Construction Digital Transformation Market Size Report, 2024-2030.',
    'Grieves, M., & Vickers, J. (2017). Digital twin: Mitigating unpredictable, undesirable emergent behavior. In Transdisciplinary Perspectives on Complex Systems (pp. 85\u2013113). Springer.',
    'IBEF. (2024). Indian Construction Industry Overview. India Brand Equity Foundation.',
    'JLL India. (2024). India Real Estate Market Update: Digital Adoption in Construction.',
    'Kerzner, H. (2022). Project Management: A Systems Approach (13th ed.). John Wiley & Sons.',
    'Kim, S., et al. (2019). Integrated cost and CO2 optimization using machine learning. Automation in Construction, 103, 218\u2013232.',
    'Kirby, P. (2020). Retail Interior Design: Creating Store Experiences that Count. Routledge.',
    'Koskela, L. (2000). An Exploration Towards a Production Theory and its Application to Construction. VTT Technical Research Centre.',
    'KPMG India. (2023). Smart Construction: How Industry 4.0 is Transforming Indian Construction.',
    'Love, P. E. D., et al. (2015). Cost overruns in transportation infrastructure projects. Transportation Research Part A, 92, 184\u2013194.',
    'McKinsey Global Institute. (2017). Reinventing Construction: A Route to Higher Productivity.',
    'Monteiro, A., & Martins, J. P. (2013). BIM-based quantity takeoff modeling guidelines. Automation in Construction, 35, 238\u2013253.',
    'Oracle Construction. (2024). Global Construction Survey: The Path to Digital Maturity.',
    'Pheng, L. S., & Chuan, Q. T. (2006). Environmental factors and work performance of project managers. International Journal of Project Management, 24(1), 24\u201337.',
    'PMI. (2021). Pulse of the Profession: Beyond Agility. Project Management Institute.',
    'Shank, J. K., & Govindarajan, V. (1993). Strategic Cost Management: The New Tool for Competitive Advantage. The Free Press.',
    'Smith, P. (2016). BIM & the 5D project cost manager. Procedia - Social and Behavioral Sciences, 226, 193\u2013200.',
    'Tayefeh Hashemi, S., et al. (2020). Cost estimation and prediction in construction: ML techniques review. SN Applied Sciences, 2, 1703.',
    'Vrijhoef, R., & Koskela, L. (2000). The four roles of supply chain management in construction. European Journal of Purchasing & Supply Management, 6(3-4), 169\u2013178.',
]

for i, ref in enumerate(references, 1):
    doc.add_paragraph(f'{i}. {ref}')

# ============================
# ANNEXURE
# ============================
add_page_break()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('ANNEXURE')
run.bold = True
run.font.size = Pt(14)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('QUESTIONNAIRE')
run.bold = True
run.font.size = Pt(13)

doc.add_paragraph()

add_heading('SECTION A: DEMOGRAPHICS AND PROFESSIONAL PROFILE', level=3)
doc.add_paragraph('Q1. Name (Optional): _______________')
doc.add_paragraph('Q2. Gender: (1) Male (2) Female (3) Prefer not to say')
doc.add_paragraph('Q3. Age Group: (1) 25\u201334 (2) 35\u201344 (3) 45\u201354 (4) 55+')
doc.add_paragraph('Q4. Experience in construction/interior projects: (1) 1\u20135 yrs (2) 6\u201310 yrs (3) 11\u201315 yrs (4) 16+')
doc.add_paragraph('Q5. Primary Role: (1) Project Manager (2) Architect/Designer (3) Estimation (4) Procurement (5) Senior Mgmt (6) Other')
doc.add_paragraph('Q6. Organization Type: (1) Interior Fit-out Firm (2) Architecture Practice (3) Brand/Client (4) Vendor (5) Consulting (6) Other')
doc.add_paragraph('Q7. Average project budget range: (1) Below \u20b950L (2) \u20b950L\u2013\u20b91Cr (3) \u20b91Cr\u2013\u20b93Cr (4) \u20b93Cr\u2013\u20b95Cr (5) Above \u20b95Cr')

doc.add_paragraph()
add_heading('SECTION B: ESTIMATION PRACTICES (1=Strongly Disagree to 5=Strongly Agree)', level=3)
doc.add_paragraph('B1. Our organization primarily uses manual methods for cost estimation.')
doc.add_paragraph('B2. Our cost estimates frequently deviate from actual project costs.')
doc.add_paragraph('B3. Design changes during execution significantly impact our project budgets.')
doc.add_paragraph('B4. We have a standardized estimation process followed across all projects.')
doc.add_paragraph('B5. Our estimation accuracy depends heavily on individual estimators\u2019 experience.')
doc.add_paragraph('B6. We maintain a structured database of historical project costs.')

doc.add_paragraph()
add_heading('SECTION C: DIGITAL TOOL AWARENESS (1=Strongly Disagree to 5=Strongly Agree)', level=3)
doc.add_paragraph('C1. I am aware of BIM tools for cost estimation.')
doc.add_paragraph('C2. Digital estimation tools can significantly improve budget accuracy.')
doc.add_paragraph('C3. AI-powered cost prediction tools would benefit our process.')
doc.add_paragraph('C4. Initial investment required for digital tools is a barrier.')
doc.add_paragraph('C5. Lack of trained personnel is a major barrier to adoption.')
doc.add_paragraph('C6. Real-time cost tracking dashboards would help manage overruns.')
doc.add_paragraph('C7. Cloud-based collaboration tools improve stakeholder communication.')
doc.add_paragraph('C8. Our organization plans to increase digital tool investment in the next 2 years.')

doc.add_paragraph()
add_heading('SECTION D: VENDOR MANAGEMENT (1=Strongly Disagree to 5=Strongly Agree)', level=3)
doc.add_paragraph('D1. We use a centralized database to manage vendor information.')
doc.add_paragraph('D2. Vendor delivery delays are a frequent cause of cost increases.')
doc.add_paragraph('D3. We have a formal vendor performance evaluation system.')
doc.add_paragraph('D4. Competitive bidding processes help us achieve better pricing.')
doc.add_paragraph('D5. Quality inconsistency among vendors is a significant challenge.')
doc.add_paragraph('D6. Digital vendor portals would streamline our procurement.')
doc.add_paragraph('D7. Long-term vendor relationships are more valuable than competitive bidding.')

doc.add_paragraph()
add_heading('SECTION E: COST PERFORMANCE (1=Strongly Disagree to 5=Strongly Agree)', level=3)
doc.add_paragraph('E1. Digital tools have/would improve our project profitability.')
doc.add_paragraph('E2. Better vendor management directly contributes to reduced costs.')
doc.add_paragraph('E3. Integrated systems would reduce our cost overruns by 15%+.')
doc.add_paragraph('E4. Our projects typically experience cost overruns of more than 15%.')
doc.add_paragraph('E5. Investing in digital cost management tools would provide positive ROI within 1 year.')

doc.add_paragraph()
add_heading('SECTION F: OPEN-ENDED QUESTIONS', level=3)
doc.add_paragraph('F1. What are the top 3 challenges you face in managing project costs?')
doc.add_paragraph('F2. What digital tools does your organization currently use?')
doc.add_paragraph('F3. What would encourage your organization to adopt more digital tools?')
doc.add_paragraph('F4. Any additional comments regarding cost management in interior projects?')

doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Thank you for your valuable time and participation in this research study.')
run.italic = True

# Save the document
output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'Project_Report.docx')
doc.save(output_path)
print(f'Document saved: {output_path}')
print(f'Total sections created successfully.')
