"""
Step 1: Regenerate the base DOCX report with expanded content.
This script creates a comprehensive 70+ page report.
Split into parts for manageability.
"""
import os
import sys
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
REPORT_DIR = os.path.join(BASE_DIR, '..')
FIGURES_DIR = os.path.join(REPORT_DIR, 'figures')
OUTPUT_PATH = os.path.join(REPORT_DIR, 'Project_Report.docx')

doc = Document()

# ============ PAGE SETUP ============
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(2.54)

style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(6)

# ============ HELPER FUNCTIONS ============
def set_cell_shading(cell, color):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h

def para(text, bold=False, italic=False, align=None, size=12, space_after=Pt(6), first_indent=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if align:
        p.alignment = align
    p.paragraph_format.space_after = space_after
    if first_indent:
        p.paragraph_format.first_line_indent = first_indent
    return p

def page_break():
    doc.add_page_break()

def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(10)
        set_cell_shading(cell, "D9E2F3")
    # Data rows
    for row_data in rows:
        row = table.add_row()
        for i, val in enumerate(row_data):
            cell = row.cells[i]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10)
    doc.add_paragraph()
    return table

def add_figure(filename, caption, width=Inches(5.5)):
    path = os.path.join(FIGURES_DIR, filename)
    if os.path.exists(path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(path, width=width)
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cap.add_run(caption)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(10)
        run.italic = True
        doc.add_paragraph()

# ============================================================
# TITLE PAGE
# ============================================================
for _ in range(4):
    doc.add_paragraph()

para("SURVEY AND ANALYSIS OF QUANTUM PROCESSING\nINTEGRATION WITH LARGE LANGUAGE MODELS (LLMs)",
     bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=16)
doc.add_paragraph()
para("Project Report Submitted in Partial Fulfilment of the\nRequirement for the Award of Degree of",
     align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
doc.add_paragraph()
para("MASTER OF BUSINESS ADMINISTRATION (MBA)", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=14)
doc.add_paragraph()
para("Submitted by", align=WD_ALIGN_PARAGRAPH.CENTER)
para("Vigneshwara Chinnadurai", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=14)
para("Reg. No.: 2414504298", align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()
para("Under the Guidance of", align=WD_ALIGN_PARAGRAPH.CENTER)
para("Mr. Govind", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=14)
for _ in range(3):
    doc.add_paragraph()
para("CENTRE FOR DISTANCE AND ONLINE EDUCATION", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
para("MANIPAL UNIVERSITY JAIPUR", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
para("May 2026", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

page_break()

# ============================================================
# BONAFIDE CERTIFICATE
# ============================================================
heading("BONAFIDE CERTIFICATE", level=1)
para("This is to certify that Vigneshwara Chinnadurai, a student of Master of Business Administration (MBA), Roll Number 2414504298, has successfully completed the project titled \"Survey and Analysis of Quantum Processing Integration with Large Language Models (LLMs)\" under my supervision as a part of the requirements for the MBA program at Centre for Distance and Online Education, Manipal University Jaipur during the academic year 2024-2026.")
para("This project report embodies the original work of the student, conducted with due diligence, and adheres to the standards expected by the institution. It has not been submitted to any other institution for any degree, diploma, or certificate.")
doc.add_paragraph()
doc.add_paragraph()
para("[Guide's Signature]", bold=True)
para("Mr. Govind", bold=True)
para("Date: May 2026")
para("Place: Bangalore, Karnataka")
page_break()

# ============================================================
# DECLARATION
# ============================================================
heading("DECLARATION BY THE STUDENT", level=1)
para("I, Vigneshwara Chinnadurai, a student of Master of Business Administration (MBA), Registration Number 2414504298, hereby declare that the project report titled \"Survey and Analysis of Quantum Processing Integration with Large Language Models (LLMs)\" submitted to Centre for Distance and Online Education, Manipal University Jaipur is a record of my original work carried out under the guidance of Mr. Govind.")
para("I affirm that this project is the result of my own independent effort, and to the best of my knowledge, it does not contain any material previously published or written by any other person or material which has been accepted for the award of any other degree or diploma at any other educational institution, except where due acknowledgment has been made in the text.")
para("I also declare that I have adhered to all the guidelines and standards required for academic honesty and have cited all sources wherever used.")
doc.add_paragraph()
doc.add_paragraph()
para("[Student's Signature]", bold=True)
para("Vigneshwara Chinnadurai", bold=True)
para("Reg. No.: 2414504298")
para("Date: May 2026")
para("Place: Bangalore, Karnataka")
page_break()

# ============================================================
# ACKNOWLEDGMENTS
# ============================================================
heading("ACKNOWLEDGMENTS", level=1)
para("I would like to express my sincere gratitude to all those who contributed to the completion of this project.")
para("First and foremost, I am profoundly grateful to my project guide, Mr. Govind, whose expert guidance, continuous encouragement, and insightful suggestions have been instrumental throughout the research process. His deep understanding of data science and emerging technologies provided a strong foundation for this study.")
para("I extend my heartfelt thanks to the Centre for Distance and Online Education, Manipal University Jaipur, for providing an excellent academic framework and resources that facilitated this research.")
para("I am also thankful to the open-source communities behind IBM Qiskit, Xanadu PennyLane, and Google Cirq for making quantum computing accessible through their simulators, documentation, and tutorials, which were critical for the experimental component of this project.")
para("I would like to acknowledge the contributions of the broader quantum computing and NLP research communities, whose published works formed the foundation of the literature review in this study.")
para("Special thanks to my colleagues and peers in the MBA Analytics and Data Science program for their intellectual discussions and moral support throughout this journey.")
para("Finally, I express my deepest appreciation to my family and friends for their unwavering support and encouragement throughout this academic journey.")
doc.add_paragraph()
para("Vigneshwara Chinnadurai", bold=True)
page_break()

# ============================================================
# TABLE OF CONTENTS
# ============================================================
heading("TABLE OF CONTENTS", level=1)
toc_items = [
    ("", "Title Page", "1"),
    ("", "Bonafide Certificate", "2"),
    ("", "Declaration", "3"),
    ("", "Acknowledgments", "4"),
    ("", "Table of Contents", "5-6"),
    ("", "List of Tables", "7"),
    ("", "List of Figures", "8"),
    ("", "List of Abbreviations", "9-10"),
    ("", "Executive Summary", "11-12"),
    ("1", "Introduction", "13"),
    ("1.1", "Background of the Study", "13"),
    ("1.2", "Statement of the Problem", "16"),
    ("1.3", "Research Objectives", "18"),
    ("1.4", "Research Questions", "19"),
    ("1.5", "Scope of the Study", "20"),
    ("1.6", "Significance of the Study", "21"),
    ("2", "Literature Review", "22"),
    ("2.1", "Evolution of Large Language Models", "22"),
    ("2.2", "Fundamentals of Quantum Computing", "25"),
    ("2.3", "Quantum Natural Language Processing", "28"),
    ("2.4", "Hybrid Quantum-Classical Architectures", "31"),
    ("2.5", "Industry Initiatives and Investments", "34"),
    ("2.6", "Research Gaps", "35"),
    ("3", "Research Methodology", "36"),
    ("3.1", "Research Design", "36"),
    ("3.2", "Data Collection Methods", "37"),
    ("3.3", "Experimental Framework", "38"),
    ("3.4", "Data Analysis Techniques", "40"),
    ("3.5", "Ethical Considerations", "41"),
    ("3.6", "Limitations of the Methodology", "42"),
    ("4", "Data Analysis and Interpretation", "43"),
    ("4.1", "Literature Analysis Results", "43"),
    ("4.2", "Experiment 1: Quantum Word Encoding", "46"),
    ("4.3", "Experiment 2: Quantum Text Classification", "49"),
    ("4.4", "Experiment 3: Hybrid Quantum-Classical NLP", "52"),
    ("4.5", "Experiment 4: Performance Benchmarking", "55"),
    ("5", "Findings and Discussion", "58"),
    ("5.1", "Key Research Findings", "58"),
    ("5.2", "Comparison with Existing Literature", "60"),
    ("5.3", "Practical Implications", "62"),
    ("6", "Conclusions", "63"),
    ("7", "Recommendations", "65"),
    ("8", "Limitations of the Study", "67"),
    ("9", "References / Bibliography", "69"),
    ("10", "Appendices", "72"),
    ("A", "Appendix A: Experimental Code Listings", "72"),
    ("B", "Appendix B: Raw Experimental Data", "78"),
    ("C", "Appendix C: Glossary of Terms", "80"),
]

table = doc.add_table(rows=0, cols=3)
table.style = 'Table Grid'
for sec, title, page in toc_items:
    row = table.add_row()
    row.cells[0].text = sec
    row.cells[1].text = title
    row.cells[2].text = page
    for cell in row.cells:
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11)
            if cell == row.cells[0]:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

page_break()

# ============================================================
# LIST OF TABLES
# ============================================================
heading("LIST OF TABLES", level=1)
tables_list = [
    ("Table 2.1", "Growth in LLM Parameters and Compute Requirements"),
    ("Table 2.2", "Current Quantum Hardware Landscape (2025)"),
    ("Table 2.3", "Key QNLP Research Timeline"),
    ("Table 2.4", "Industry Quantum Computing Investments"),
    ("Table 3.1", "Experimental Configuration Summary"),
    ("Table 4.1", "Literature Distribution by Approach Type"),
    ("Table 4.2", "Literature Distribution by Research Focus"),
    ("Table 4.3", "Technology Readiness Level Assessment"),
    ("Table 4.4", "Encoding Efficiency Comparison"),
    ("Table 4.5", "Semantic Preservation Results"),
    ("Table 4.6", "Classification Performance Comparison"),
    ("Table 4.7", "Training Convergence Analysis"),
    ("Table 4.8", "Pipeline Comparison Results"),
    ("Table 4.9", "Scalability Analysis"),
    ("Table 4.10", "Noise Impact on Quantum Classifier"),
    ("Table 4.11", "Cross-Framework Comparison"),
    ("Table 4.12", "Summary Comparison Matrix"),
    ("Table B.1", "Raw Encoding Fidelity Data"),
    ("Table B.2", "Raw Classification Results per Fold"),
    ("Table B.3", "Raw Noise Analysis Data"),
]
for tnum, title in tables_list:
    p = doc.add_paragraph()
    run = p.add_run(f"{tnum}: {title}")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)
page_break()

# ============================================================
# LIST OF FIGURES
# ============================================================
heading("LIST OF FIGURES", level=1)
figures_list = [
    ("Figure 4.1", "Classical Cosine Similarity Matrix of Word Embeddings"),
    ("Figure 4.2", "Semantic Preservation Across Encoding Methods"),
    ("Figure 4.3", "Experiment 1 Comparative Results"),
    ("Figure 4.4", "Quantum Text Classification Results"),
    ("Figure 4.5", "Hybrid Pipeline Learning Curves and Parameter Efficiency"),
    ("Figure 4.6", "Noise Resilience and Evaluation Matrix"),
    ("Figure 4.7", "Technology Maturity Roadmap"),
    ("Figure 4.8", "Strategic Adoption Framework"),
]
for fnum, title in figures_list:
    p = doc.add_paragraph()
    run = p.add_run(f"{fnum}: {title}")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)
page_break()

# ============================================================
# LIST OF ABBREVIATIONS
# ============================================================
heading("LIST OF ABBREVIATIONS", level=1)
abbreviations = [
    ("AI", "Artificial Intelligence"),
    ("API", "Application Programming Interface"),
    ("BERT", "Bidirectional Encoder Representations from Transformers"),
    ("CNOT", "Controlled-NOT (quantum gate)"),
    ("DisCoCat", "Distributional Compositional Categorical"),
    ("DNN", "Deep Neural Network"),
    ("F1", "F1 Score (harmonic mean of precision and recall)"),
    ("GloVe", "Global Vectors for Word Representation"),
    ("GPT", "Generative Pre-trained Transformer"),
    ("GPU", "Graphics Processing Unit"),
    ("IBM", "International Business Machines Corporation"),
    ("IQP", "Instantaneous Quantum Polynomial"),
    ("LLM", "Large Language Model"),
    ("LSTM", "Long Short-Term Memory"),
    ("ML", "Machine Learning"),
    ("MLP", "Multi-Layer Perceptron"),
    ("MoE", "Mixture of Experts"),
    ("NISQ", "Noisy Intermediate-Scale Quantum"),
    ("NLP", "Natural Language Processing"),
    ("NN", "Neural Network"),
    ("PCA", "Principal Component Analysis"),
    ("PF-days", "Petaflop-days (unit of computational work)"),
    ("QAOA", "Quantum Approximate Optimization Algorithm"),
    ("QC", "Quantum Computing / Quantum Classifier"),
    ("QML", "Quantum Machine Learning"),
    ("QNLP", "Quantum Natural Language Processing"),
    ("QPU", "Quantum Processing Unit"),
    ("RNN", "Recurrent Neural Network"),
    ("SVM", "Support Vector Machine"),
    ("TF-IDF", "Term Frequency-Inverse Document Frequency"),
    ("TRL", "Technology Readiness Level"),
    ("VQC", "Variational Quantum Circuit"),
    ("VQE", "Variational Quantum Eigensolver"),
    ("Word2Vec", "Word to Vector (word embedding model)"),
]
add_table(["Abbreviation", "Full Form"], [(a, f) for a, f in abbreviations])
page_break()

print("[Part 1] Front matter complete. Saving intermediate...")
doc.save(OUTPUT_PATH)
print(f"  Saved to: {OUTPUT_PATH}")
