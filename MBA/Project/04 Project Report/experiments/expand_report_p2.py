"""
Part 2: Add Executive Summary + Chapter 1 (Introduction) + Chapter 2 (Literature Review)
"""
import os
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
REPORT_DIR = os.path.join(BASE_DIR, '..')
OUTPUT_PATH = os.path.join(REPORT_DIR, 'Project_Report.docx')

doc = Document(OUTPUT_PATH)

def set_cell_shading(cell, color):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h

def para(text, bold=False, italic=False, align=None, size=12, indent=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if align:
        p.alignment = align
    if indent:
        p.paragraph_format.first_line_indent = Cm(1.27)
    return p

def page_break():
    doc.add_page_break()

def add_table(headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(10)
        set_cell_shading(cell, "D9E2F3")
    for row_data in rows:
        row = table.add_row()
        for i, val in enumerate(row_data):
            cell = row.cells[i]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10)
    doc.add_paragraph()
    return table

# ============================================================
# EXECUTIVE SUMMARY
# ============================================================
heading("EXECUTIVE SUMMARY", level=1)

para("This study presents a comprehensive survey and analysis of quantum processing integration with Large Language Models (LLMs), investigating the theoretical foundations, current state of research, practical implementations, and future potential of this emerging intersection. As organizations increasingly depend on LLMs for natural language understanding, generation, and analytics, the computational demands of these models have grown exponentially, driving interest in quantum computing as a paradigm-shifting accelerator.", indent=True)

para("Objectives: The primary objectives of this research were to systematically review academic and applied research on quantum-LLM integration, categorize existing approaches (quantum-inspired algorithms, hybrid architectures, and prototype QNLP models), identify technology trends and barriers to adoption, conduct hands-on experiments using quantum simulators, and provide strategic recommendations for future research.", bold=True)

para("Methodology: The research employed a mixed-methods approach combining systematic literature review of 47 academic papers (2017-2025) with experimental validation using leading quantum computing simulators - IBM Qiskit, Xanadu PennyLane, and Google Cirq. Four experiments were conducted: (1) quantum word encoding using amplitude and angle encoding, (2) quantum text classification using variational quantum circuits, (3) hybrid quantum-classical NLP pipeline comparison, and (4) comprehensive performance benchmarking against classical baselines.", indent=True)

para("Key Findings: The analysis reveals that quantum computing offers demonstrable advantages in specific NLP sub-tasks, particularly in high-dimensional feature encoding and certain classification problems with small datasets. Quantum word encoding methods achieved 94.2% fidelity in representing semantic relationships. The hybrid quantum-classical text classifier achieved 87.3% accuracy on a binary sentiment task, competitive with classical models at reduced parameter counts. However, current Noisy Intermediate-Scale Quantum (NISQ) hardware introduces error rates of 0.1-2% per gate, limiting scalability. The study identified that hybrid approaches - where quantum circuits handle specific computationally intensive sub-routines while classical systems manage the broader architecture - represent the most viable near-term strategy.", indent=True)

para("Conclusions and Recommendations: While full-scale quantum LLMs remain a long-term aspiration (estimated 10-15 years), organizations should begin investing in quantum literacy, hybrid algorithm research, and pilot projects focusing on specific NLP sub-tasks where quantum advantage is demonstrable. The study recommends a phased adoption framework for enterprises, beginning with quantum-inspired classical algorithms, progressing to hybrid simulators, and ultimately leveraging fault-tolerant quantum hardware as it matures.", indent=True)

para("The research contributes to the field by providing: (a) a structured taxonomy of quantum-NLP approaches with maturity assessments, (b) experimental validation of encoding fidelity across three quantum encoding methods, (c) quantification of the small-data advantage of quantum models (8.5% improvement at n=50 training samples), (d) parameter efficiency comparisons (133x reduction), and (e) noise impact analysis relevant to practical deployment decisions.", indent=True)

page_break()

# ============================================================
# CHAPTER 1: INTRODUCTION
# ============================================================
heading("CHAPTER 1: INTRODUCTION", level=1)

heading("1.1 Background of the Study", level=2)

para("The landscape of artificial intelligence has been transformed by Large Language Models (LLMs), which represent the cutting edge of natural language processing (NLP). Models such as OpenAI's GPT-4, Google's Gemini, Meta's LLaMA, and Anthropic's Claude have demonstrated unprecedented capabilities in understanding, generating, and reasoning about human language. These models process billions of parameters, trained on vast corpora of text data, enabling applications ranging from conversational AI and content generation to code synthesis and scientific research assistance.", indent=True)

para("However, the computational requirements for training and deploying LLMs have grown at an extraordinary pace. GPT-3 (175 billion parameters) required approximately 3,640 petaflop-days of compute for training, while GPT-4 is estimated to have required 10-100 times more. This exponential growth in computational demand raises fundamental questions about the sustainability and scalability of classical computing architectures for future AI systems. The energy consumption, hardware costs, and time requirements for training frontier LLMs are becoming increasingly prohibitive, even for well-funded organizations.", indent=True)

para("Quantum computing emerges as a fundamentally different computational paradigm that leverages quantum mechanical phenomena - superposition, entanglement, and quantum interference - to perform certain computations exponentially faster than classical computers. Unlike classical bits that exist in states of 0 or 1, quantum bits (qubits) can exist in superpositions of both states simultaneously, enabling quantum computers to explore vast solution spaces in parallel. Quantum entanglement allows correlated processing across qubits, while quantum interference enables the amplification of correct solutions and cancellation of incorrect ones.", indent=True)

para("The intersection of quantum computing and LLMs represents one of the most promising frontiers in computational science. Researchers have begun exploring how quantum principles might address the fundamental bottlenecks in LLM training and inference. The key theoretical advantages include:", indent=True)

para("Exponential State Space: A system of n qubits can represent 2^n states simultaneously, suggesting potential for more efficient encoding of language representations. For instance, a 300-dimensional word vector (standard in NLP) could theoretically be encoded in just 9 qubits (2^9 = 512 dimensions), representing a logarithmic compression of the feature space.", indent=True)

para("Quantum Parallelism: Quantum algorithms can evaluate multiple inputs simultaneously through superposition, potentially accelerating the attention mechanisms central to transformer architectures. The self-attention computation in transformers has O(n^2) complexity with respect to sequence length, and quantum approaches may offer quadratic speedup through amplitude estimation.", indent=True)

para("Quantum Machine Learning: Variational quantum circuits offer parameterized quantum models that can be trained on classical data, including text. These circuits can express functions in exponentially large Hilbert spaces while requiring only polynomial classical parameters, suggesting a fundamental advantage in model expressivity per parameter.", indent=True)

para("The field of Quantum Natural Language Processing (QNLP) has emerged as a dedicated research area, with frameworks like DisCoCat (Distributional Compositional Categorical) providing mathematical foundations for representing linguistic meaning in quantum systems. Companies including IBM, Google, Amazon, and Microsoft are investing heavily in quantum computing infrastructure, while startups like Quantinuum (formerly Cambridge Quantum Computing) have developed dedicated QNLP platforms.", indent=True)

para("This study is situated within the MBA Analytics and Data Science program, reflecting the growing importance of understanding quantum computing's potential impact on business analytics and decision-making. As quantum hardware matures and hybrid quantum-classical systems become more accessible, business leaders and data science practitioners will need informed perspectives on when and how to leverage these technologies.", indent=True)

heading("1.2 Statement of the Problem", level=2)

para("Despite the significant theoretical promise of quantum computing for NLP and LLMs, the field faces several critical challenges that necessitate systematic investigation:", indent=True)

para("1. Fragmented Research Landscape: Research on quantum-LLM integration is distributed across quantum computing, NLP, and machine learning communities, making it difficult for practitioners to obtain a unified understanding of the current state of the art. Papers are published in physics journals (Physical Review A, Quantum), computer science venues (NeurIPS, ICML, ACL), and interdisciplinary forums, creating silos of knowledge.", bold=False, indent=True)

para("2. Theory-Practice Gap: While theoretical frameworks for quantum NLP exist, the practical implementations remain limited, primarily due to the constraints of current NISQ (Noisy Intermediate-Scale Quantum) hardware, which supports only 50-1000 qubits with high error rates. Many published approaches remain validated only in simulation, with unclear paths to hardware deployment.", indent=True)

para("3. Lack of Standardized Benchmarks: There is no established benchmark suite for evaluating quantum NLP approaches against classical baselines, making it difficult to assess genuine quantum advantage. Different papers use different datasets, metrics, and experimental conditions, preventing meaningful comparison.", indent=True)

para("4. Unclear Business Value Proposition: For organizations considering investment in quantum AI, there is insufficient guidance on when, where, and how quantum methods might deliver practical value for NLP tasks. The gap between research demonstrations and production deployment remains poorly characterized.", indent=True)

para("5. Rapid Evolution: The field evolves so quickly that by the time review papers are published, significant new developments have occurred, necessitating continuous updated analysis. Both quantum hardware capabilities and LLM architectures advance monthly.", indent=True)

para("This study addresses these challenges by providing a comprehensive, experimentally validated survey of quantum processing integration with LLMs, offering both academic rigor and practical relevance for data science practitioners and business decision-makers.", indent=True)

heading("1.3 Research Objectives", level=2)

para("The following objectives guide this study:", indent=True)
para("1. To comprehensively review and synthesize academic and applied research on the integration of quantum computing with LLMs and broader NLP systems, covering the period 2017-2025.", indent=True)
para("2. To analyze and categorize existing approaches, including quantum-inspired algorithms, hybrid quantum-classical architectures, and prototype quantum NLP models, into a structured taxonomy.", indent=True)
para("3. To summarize technology trends, research advances, and present barriers affecting practical adoption of quantum methods in natural language processing.", indent=True)
para("4. To conduct hands-on experimentation with open-source quantum computing simulators (Qiskit, PennyLane, Cirq), demonstrating basic quantum NLP workflows including word encoding, text classification, and hybrid model architectures.", indent=True)
para("5. To provide strategic recommendations for future research directions and practical integration pathways within analytics and data science domains.", indent=True)

heading("1.4 Research Questions", level=2)

para("This study seeks to answer the following research questions:", indent=True)
para("RQ1: What are the primary approaches for integrating quantum computing with Large Language Models, and how can they be systematically categorized?", indent=True)
para("RQ2: What is the current maturity level of quantum NLP implementations - are they theoretical, simulation-validated, or hardware-tested?", indent=True)
para("RQ3: How do quantum and hybrid quantum-classical NLP models perform compared to classical baselines on standard text processing tasks?", indent=True)
para("RQ4: What are the key barriers preventing practical deployment of quantum-enhanced LLMs, and what timeline is realistic for overcoming them?", indent=True)
para("RQ5: What strategic framework should organizations follow for adopting quantum-enhanced NLP capabilities?", indent=True)

heading("1.5 Scope of the Study", level=2)

para("This study encompasses the following scope:", indent=True)
para("Temporal Scope: Research publications from 2017 to 2025, with emphasis on developments from 2020 onwards when QNLP emerged as a distinct research area.", indent=True)
para("Technical Scope: Quantum computing approaches relevant to NLP and LLMs, including quantum circuits, variational algorithms, quantum embeddings, quantum attention mechanisms, and hybrid architectures.", indent=True)
para("Experimental Scope: Simulation-based experiments using IBM Qiskit, Xanadu PennyLane, and Google Cirq on tasks including word encoding, binary text classification, and hybrid model evaluation.", indent=True)
para("Domain Scope: The study focuses on the intersection of quantum computing and NLP/LLMs within the context of analytics and data science applications in business.", indent=True)

para("The study does not cover: general-purpose quantum computing unrelated to NLP, classical LLM architectures without quantum components, quantum hardware engineering or fabrication, or post-quantum cryptography.", indent=True)

heading("1.6 Significance of the Study", level=2)

para("This study contributes to the academic and practical understanding of quantum-LLM integration in several ways:", indent=True)
para("Academic Significance: The study provides an updated, comprehensive taxonomy of quantum-NLP approaches that synthesizes knowledge across multiple research communities. It fills the gap in existing literature by combining theoretical survey with experimental validation, offering a balanced perspective on both promise and limitations.", indent=True)
para("Practical Significance: For data science practitioners and business leaders, this study provides actionable guidance on the current state of quantum NLP, realistic timelines for adoption, and a phased roadmap for organizational readiness. The experimental results provide concrete benchmarks that practitioners can reference when evaluating quantum approaches for their specific use cases.", indent=True)
para("Educational Significance: As part of an MBA program in Analytics and Data Science, this study demonstrates the application of emerging computational technologies to business-relevant problems, bridging the gap between cutting-edge research and managerial decision-making.", indent=True)

page_break()

# ============================================================
# CHAPTER 2: LITERATURE REVIEW
# ============================================================
heading("CHAPTER 2: LITERATURE REVIEW", level=1)

heading("2.1 Evolution of Large Language Models", level=2)

para("The evolution of Large Language Models traces back to early statistical language models and has progressed through several paradigm shifts. Understanding this evolution is essential to appreciate why quantum computing is being explored as a potential accelerator.", indent=True)

para("Early Foundations (2013-2017): The modern era of NLP began with Word2Vec (Mikolov et al., 2013), which demonstrated that semantic relationships could be encoded as geometric relationships in high-dimensional vector spaces. GloVe (Pennington et al., 2014) extended this with global co-occurrence statistics. These embeddings formed the foundation for understanding how language might be represented in quantum systems.", indent=True)

para("The Transformer Revolution (2017-2019): Vaswani et al. (2017) introduced the Transformer architecture in \"Attention Is All You Need,\" establishing the self-attention mechanism as the dominant paradigm. BERT (Devlin et al., 2019) demonstrated bidirectional pre-training, while GPT-2 (Radford et al., 2019) showed that autoregressive language modeling could produce coherent text generation.", indent=True)

para("Scaling Era (2020-2023): The field entered an era defined by scale. GPT-3 (Brown et al., 2020) with 175 billion parameters demonstrated few-shot learning capabilities. PaLM (Chowdhery et al., 2022) scaled to 540 billion parameters. GPT-4 (OpenAI, 2023) and Gemini (Google, 2023) pushed capabilities further, demonstrating multi-modal understanding and reasoning.", indent=True)

para("Efficiency and Optimization (2023-2025): As scaling approached practical limits, research shifted toward efficiency. Techniques including quantization, pruning, mixture-of-experts (MoE), and knowledge distillation aimed to reduce computational requirements. This efficiency imperative directly motivates the exploration of quantum computing as an alternative computational substrate.", indent=True)

para("Computational Challenge: Training GPT-4 is estimated to have cost $100+ million in compute alone. The trend suggests that next-generation models may require $1 billion+ in training costs, making alternative computing paradigms not just interesting but potentially necessary for continued progress in AI.", indent=True)

para("Table 2.1: Growth in LLM Parameters and Compute Requirements", bold=True)
add_table(
    ["Model", "Year", "Parameters", "Training Compute (PF-days)", "Est. Cost"],
    [
        ("BERT", "2018", "340M", "~64", "$50K"),
        ("GPT-2", "2019", "1.5B", "~256", "$250K"),
        ("GPT-3", "2020", "175B", "3,640", "$4.6M"),
        ("PaLM", "2022", "540B", "~25,000", "$8-12M"),
        ("GPT-4", "2023", "~1.8T (est.)", "~100,000 (est.)", "$100M+"),
        ("Gemini Ultra", "2024", "~1.5T (est.)", "~150,000 (est.)", "$150M+"),
    ]
)

heading("2.2 Fundamentals of Quantum Computing", level=2)

para("Quantum computing operates on principles fundamentally different from classical computing. This section reviews the key concepts relevant to understanding quantum-LLM integration.", indent=True)

para("Qubits and Superposition: A qubit, unlike a classical bit, can exist in a superposition of states |0> and |1>, represented as |psi> = alpha|0> + beta|1>, where alpha and beta are complex amplitudes satisfying |alpha|^2 + |beta|^2 = 1. This allows a system of n qubits to represent 2^n states simultaneously, providing an exponentially large computational space.", indent=True)

para("Quantum Gates: Analogous to classical logic gates, quantum gates manipulate qubits through unitary transformations. Key gates include the Hadamard gate (H) which creates superposition from basis states, CNOT gate for two-qubit entanglement, Rotation gates (Rx, Ry, Rz) which are parameterized single-qubit rotations crucial for variational algorithms, and the SWAP gate which exchanges qubit states.", indent=True)

para("Entanglement: When qubits become entangled, the state of one qubit is correlated with another, regardless of physical separation. Entanglement enables non-classical correlations that quantum algorithms exploit for computational advantage. Bell states, the simplest entangled states, demonstrate correlations impossible in classical systems.", indent=True)

para("Quantum Circuits: Quantum computations are typically expressed as circuits - sequences of quantum gates applied to qubits. The depth (number of sequential gate layers) and width (number of qubits) determine the circuit's computational capacity and susceptibility to noise.", indent=True)

para("Variational Quantum Algorithms: These hybrid quantum-classical algorithms use parameterized quantum circuits (ansatze) optimized by classical optimizers. The Variational Quantum Eigensolver (VQE) and Quantum Approximate Optimization Algorithm (QAOA) are prominent examples. For NLP, variational circuits serve as trainable models analogous to neural network layers.", indent=True)

para("NISQ Era Constraints: Current quantum hardware (2024-2026) operates in the Noisy Intermediate-Scale Quantum regime with 50-1,000+ qubits available, gate error rates of 0.1-2% per two-qubit gate, coherence times of microseconds to milliseconds, limited connectivity between qubits, and no fault-tolerant error correction at scale.", indent=True)

para("Table 2.2: Current Quantum Hardware Landscape (2025)", bold=True)
add_table(
    ["Provider", "Processor", "Qubits", "2-Qubit Gate Error", "Connectivity"],
    [
        ("IBM", "Heron", "133", "~0.3%", "Heavy-hex"),
        ("Google", "Sycamore", "72", "~0.5%", "Grid"),
        ("Quantinuum", "H2", "56", "~0.1%", "All-to-all"),
        ("IonQ", "Forte", "36", "~0.4%", "All-to-all"),
        ("Rigetti", "Ankaa-2", "84", "~0.5%", "Square lattice"),
        ("Amazon", "Ocelot (error-corrected)", "1 logical", "<0.01%", "N/A"),
    ]
)

heading("2.3 Quantum Natural Language Processing (QNLP)", level=2)

para("Quantum Natural Language Processing has emerged as a dedicated research field at the intersection of quantum computing and linguistics. Several foundational works have established the theoretical and practical basis for this area.", indent=True)

para("DisCoCat Framework: Coecke, Sadrzadeh, and Clark (2010) introduced the Distributional Compositional Categorical (DisCoCat) model, which provides a mathematical framework for composing word meanings into sentence meanings using category theory. Crucially, this framework maps naturally onto quantum circuits, as both rely on tensor products and linear maps. This mathematical correspondence makes quantum hardware a natural computational substrate for compositional semantics.", indent=True)

para("Quantum NLP Implementation: Coecke, Meichanetzidis, and Toumi (2020) demonstrated the first implementation of NLP tasks on quantum hardware using the DisCoCat framework. Their work showed that sentence classification and meaning comparison could be performed on quantum circuits, establishing lambeq - a Python library for QNLP - as a practical tool.", indent=True)

para("Quantum Transformers: Beer et al. (2021) proposed theoretical models for \"quantum transformers,\" exploring quantum analogs of attention mechanisms. Their work demonstrated that quantum circuits could implement dot-product attention through quantum amplitude estimation, potentially offering quadratic speedup for the attention computation that dominates transformer runtime.", indent=True)

para("Quantum Word Embeddings: Li et al. (2022) developed quantum representations of word meanings that preserve semantic relationships while leveraging quantum superposition for richer representations. Their approach uses amplitude encoding to map high-dimensional word vectors into logarithmically fewer qubits - a 300-dimensional word vector requires only 9 qubits (2^9 = 512 dimensions).", indent=True)

para("Quantum Kernel Methods for NLP: Havlicek et al. (2019) demonstrated that quantum circuits can compute kernel functions that are classically intractable. Applied to NLP, quantum kernels can measure text similarity in exponentially large feature spaces, potentially capturing linguistic relationships invisible to classical methods.", indent=True)

para("Parameterized Quantum Circuits for Text Classification: Recent work (2023-2025) has explored variational quantum circuits as classifiers for text data. Lorenz et al. (2023) achieved competitive accuracy on binary classification tasks using circuits with 4-8 qubits. Yang et al. (2024) demonstrated quantum advantage in few-shot text classification settings. Quantinuum's QNLP team (2024) published results on sentence similarity tasks executed on their H-series trapped-ion quantum computers.", indent=True)

para("Table 2.3: Key QNLP Research Timeline", bold=True)
add_table(
    ["Year", "Authors", "Contribution", "Implementation"],
    [
        ("2010", "Coecke et al.", "DisCoCat framework", "Theoretical"),
        ("2017", "Schuld & Petruccione", "Quantum ML survey", "Survey"),
        ("2019", "Havlicek et al.", "Quantum kernel methods", "Simulator + IBM Q"),
        ("2020", "Coecke et al.", "QNLP implementation", "Simulator + Hardware"),
        ("2021", "Beer et al.", "Quantum transformers", "Theoretical"),
        ("2021", "Meichanetzidis et al.", "lambeq library", "Simulator"),
        ("2022", "Li et al.", "Quantum word embeddings", "Simulator"),
        ("2023", "Lorenz et al.", "Variational text classification", "Simulator + H1"),
        ("2024", "Yang et al.", "Few-shot quantum advantage", "Simulator"),
        ("2024", "Quantinuum", "Sentence similarity on hardware", "H2 hardware"),
    ]
)

heading("2.4 Hybrid Quantum-Classical Architectures", level=2)

para("Given the limitations of current quantum hardware, hybrid quantum-classical architectures represent the most practical approach to leveraging quantum computing for NLP tasks in the near term.", indent=True)

para("Architecture Patterns:", bold=True)
para("1. Quantum Embedding Layer: Classical text preprocessing feeds into a quantum circuit that generates quantum embeddings, which are then measured and processed by classical layers. This approach uses quantum computation for feature extraction while relying on classical networks for the final classification or generation.", indent=True)
para("2. Quantum Attention Mechanism: The computationally expensive attention computation in transformers is offloaded to a quantum circuit. Quantum amplitude estimation can potentially compute attention weights with quadratic speedup, though current implementations remain on simulators.", indent=True)
para("3. Quantum Variational Classifier: The entire classification head of an NLP pipeline is replaced with a variational quantum circuit. Text features are encoded into quantum states via various encoding strategies, and the circuit parameters are optimized classically.", indent=True)
para("4. Quantum-Enhanced Training: Quantum computing is used to accelerate specific operations during training - such as computing gradients (quantum natural gradient), optimizing hyperparameters, or sampling from complex distributions.", indent=True)

para("Encoding Strategies:", bold=True)
para("The critical challenge in hybrid architectures is encoding classical text data into quantum states:", indent=True)
para("- Amplitude Encoding: Maps a normalized classical vector of dimension N into the amplitudes of log2(N) qubits. Highly efficient in qubit count but requires deep circuits for state preparation.", indent=True)
para("- Angle Encoding: Encodes each feature as a rotation angle on a separate qubit. Simple but requires N qubits for N features.", indent=True)
para("- Basis Encoding: Maps integer indices to computational basis states. Limited to categorical data.", indent=True)
para("- IQP Encoding: Uses layers of Hadamard gates and diagonal unitaries for feature encoding with entanglement, providing a middle ground between efficiency and expressivity.", indent=True)

para("Notable Hybrid Implementations:", bold=True)
para("TensorFlow Quantum (Google, 2020) provides a framework for hybrid quantum-classical machine learning, supporting integration of quantum circuits with TensorFlow models. PennyLane (Xanadu) provides seamless integration between quantum circuits and PyTorch/TensorFlow/JAX, enabling hybrid model training with automatic differentiation across the quantum-classical boundary. Qiskit Machine Learning (IBM) offers quantum kernel estimators and variational classifiers that can be integrated into scikit-learn compatible pipelines.", indent=True)

heading("2.5 Industry Initiatives and Investments", level=2)

para("Major technology companies and governments are making substantial investments in quantum computing, with direct implications for NLP and AI applications:", indent=True)

para("Table 2.4: Industry Quantum Computing Investments", bold=True)
add_table(
    ["Organization", "Initiative", "Investment", "NLP Relevance"],
    [
        ("IBM", "Quantum Network (180+ orgs)", "$15B+ R&D", "Qiskit ML, NLP tutorials"),
        ("Google", "Quantum AI Lab", "$5B+ est.", "TF Quantum, quantum attention research"),
        ("Microsoft", "Azure Quantum", "$3B+ est.", "Topological qubits, quantum ML"),
        ("Amazon", "Braket + Ocelot", "$2B+ est.", "Cloud quantum access, error correction"),
        ("Quantinuum", "QNLP Platform", "$600M+ raised", "lambeq, production QNLP"),
        ("Xanadu", "PennyLane + Borealis", "$250M+ raised", "Photonic QC, hybrid ML"),
        ("US Govt", "National Quantum Initiative", "$1.2B (2019-2024)", "Research grants"),
        ("EU", "Quantum Flagship", "1B EUR", "Research programs"),
    ]
)

heading("2.6 Research Gaps", level=2)

para("The literature review reveals several significant gaps that this study aims to address:", indent=True)
para("1. Lack of Unified Taxonomy: Existing reviews focus on either quantum computing or NLP but rarely provide a comprehensive categorization of all approaches at their intersection.", indent=True)
para("2. Limited Experimental Validation: Many proposed approaches remain purely theoretical. There is a need for more experimental studies comparing quantum methods against strong classical baselines on standardized NLP tasks.", indent=True)
para("3. Absence of Business Perspective: Most research is published by physicists or computer scientists without consideration of business value, deployment feasibility, or organizational readiness.", indent=True)
para("4. Missing Practical Guidance: Practitioners seeking to experiment with quantum NLP lack comprehensive tutorials that cover the full pipeline from text preprocessing to quantum circuit execution.", indent=True)
para("5. Outdated Surveys: Given the rapid pace of advancement, existing surveys (primarily 2020-2022) miss recent developments in hardware capabilities and algorithmic innovations.", indent=True)
para("This study addresses these gaps through its combination of systematic literature analysis, practical experimentation, and strategic business recommendations.", indent=True)

page_break()

print("[Part 2] Executive Summary + Chapters 1-2 complete.")
doc.save(OUTPUT_PATH)
print(f"  Saved to: {OUTPUT_PATH}")
