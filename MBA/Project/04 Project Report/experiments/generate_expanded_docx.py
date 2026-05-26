"""
Generate expanded 70+ page DOCX report.
Author: Vigneshwara Chinnadurai (2414504298)
"""
import os
import sys
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import copy

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
REPORT_DIR = os.path.join(BASE_DIR, '..')
FIGURES_DIR = os.path.join(REPORT_DIR, 'figures')
OUTPUT_PATH = os.path.join(REPORT_DIR, 'Project_Report.docx')

doc = Document()

# ============ PAGE SETUP ============
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(2.54)

style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5

# Helper functions
def set_cell_shading(cell, color):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h

def add_para(text, bold=False, italic=False, alignment=None, space_after=Pt(6)):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.bold = bold
    run.italic = italic
    if alignment:
        p.alignment = alignment
    p.paragraph_format.space_after = space_after
    return p

def add_table_with_data(headers, rows, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(10)
        set_cell_shading(cell, "D9E2F3")
    # Data rows
    for row_data in rows:
        row = table.add_row()
        for i, val in enumerate(row_data):
            cell = row.cells[i]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10)
    doc.add_paragraph()  # spacing
    return table

def add_code_block(code_text):
    """Add formatted code block."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(code_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    return p

def add_figure(image_path, caption, width=Inches(5.5)):
    """Add figure with caption."""
    if os.path.exists(image_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(image_path, width=width)
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cap.add_run(caption)
        run.font.size = Pt(10)
        run.italic = True
        cap.paragraph_format.space_after = Pt(12)

def page_break():
    doc.add_page_break()

# ============================================================
# TITLE PAGE
# ============================================================
for _ in range(4):
    doc.add_paragraph()

add_para("SURVEY AND ANALYSIS OF QUANTUM PROCESSING INTEGRATION WITH LARGE LANGUAGE MODELS (LLMs)",
         bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(24))
add_para("Project Report Submitted in Partial Fulfilment of the Requirement for the Award of Degree of",
         alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(12))
add_para("MASTER OF BUSINESS ADMINISTRATION (MBA)", bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(24))
add_para("Specialization: Analytics and Data Science", alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(36))
add_para("Submitted by", alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(12))
add_para("VIGNESHWARA CHINNADURAI", bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(6))
add_para("Registration No.: 2414504298", alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(36))
add_para("Under the Guidance of", alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(12))
add_para("Mr. Govind", bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(48))
add_para("CENTRE FOR DISTANCE AND ONLINE EDUCATION", bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(6))
add_para("MANIPAL UNIVERSITY JAIPUR", bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(12))
add_para("May 2026", alignment=WD_ALIGN_PARAGRAPH.CENTER)

page_break()

# ============================================================
# BONAFIDE CERTIFICATE
# ============================================================
add_heading_styled("BONAFIDE CERTIFICATE", level=1)
doc.add_paragraph()
add_para("This is to certify that Vigneshwara Chinnadurai, a student of Master of Business Administration (MBA), Registration Number 2414504298, has successfully completed the project titled \"Survey and Analysis of Quantum Processing Integration with Large Language Models (LLMs)\" under my supervision as a part of the requirements for the MBA program at Centre for Distance and Online Education, Manipal University Jaipur during the academic year 2024\u20132026.")
doc.add_paragraph()
add_para("This project report embodies the original work of the student, conducted with due diligence, and adheres to the standards expected by the institution. It has not been submitted to any other institution for any degree, diploma, or certificate.")
doc.add_paragraph()
add_para("The project demonstrates a thorough understanding of the subject matter, including quantum computing fundamentals, natural language processing, and their intersection. The experimental work conducted using quantum simulators shows independent research capability and analytical skills appropriate for an MBA graduate.")
doc.add_paragraph()
doc.add_paragraph()
add_para("[Guide's Signature]", bold=True)
add_para("Mr. Govind", bold=True)
add_para("Project Guide")
doc.add_paragraph()
add_para("Date: May 2026")
add_para("Place: Bangalore, Karnataka")

page_break()

# ============================================================
# DECLARATION
# ============================================================
add_heading_styled("DECLARATION BY THE STUDENT", level=1)
doc.add_paragraph()
add_para("I, Vigneshwara Chinnadurai, a student of Master of Business Administration (MBA), Registration Number 2414504298, hereby declare that the project report titled \"Survey and Analysis of Quantum Processing Integration with Large Language Models (LLMs)\" submitted to Centre for Distance and Online Education, Manipal University Jaipur is a record of my original work carried out under the guidance of Mr. Govind.")
doc.add_paragraph()
add_para("I affirm that this project is the result of my own independent effort, and to the best of my knowledge, it does not contain any material previously published or written by any other person or material which has been accepted for the award of any other degree or diploma at any other educational institution, except where due acknowledgment has been made in the text.")
doc.add_paragraph()
add_para("I also declare that I have adhered to all the guidelines and standards required for academic honesty and have cited all sources wherever used. The data presented in this report is genuine and has been collected and analyzed as per the methodology described herein.")
doc.add_paragraph()
add_para("I understand that any false declaration or misrepresentation of facts will make me liable for disciplinary action as per the university rules.")
doc.add_paragraph()
doc.add_paragraph()
add_para("[Student's Signature]", bold=True)
add_para("Vigneshwara Chinnadurai", bold=True)
add_para("Reg. No.: 2414504298")
doc.add_paragraph()
add_para("Date: May 2026")
add_para("Place: Bangalore, Karnataka")

page_break()

# ============================================================
# ACKNOWLEDGMENTS
# ============================================================
add_heading_styled("ACKNOWLEDGMENTS", level=1)
doc.add_paragraph()
add_para("I would like to express my sincere gratitude to all those who contributed to the successful completion of this project report.")
doc.add_paragraph()
add_para("First and foremost, I am profoundly grateful to my project guide, Mr. Govind, whose expert guidance, continuous encouragement, and insightful suggestions have been instrumental throughout the research process. His deep understanding of data science and emerging technologies provided a strong foundation for this study. His patience in reviewing multiple drafts and providing constructive feedback has significantly improved the quality of this work.")
doc.add_paragraph()
add_para("I extend my heartfelt thanks to the Centre for Distance and Online Education, Manipal University Jaipur, for providing an excellent academic framework and resources that facilitated this research. The flexibility of the distance learning program allowed me to balance professional responsibilities with academic pursuits effectively.")
doc.add_paragraph()
add_para("I am also thankful to the open-source communities behind IBM Qiskit, Xanadu PennyLane, and Google Cirq for making quantum computing accessible through their simulators, documentation, and tutorials, which were critical for the experimental component of this project. The availability of these free tools democratizes access to quantum computing research.")
doc.add_paragraph()
add_para("I would like to acknowledge the researchers whose published work formed the foundation of my literature review. The quantum computing and NLP research communities have been remarkably open in sharing their findings, code, and datasets, which greatly facilitated this study.")
doc.add_paragraph()
add_para("Finally, I express my deepest appreciation to my family and friends for their unwavering support and encouragement throughout this academic journey. Their understanding during the long hours of research and writing made this achievement possible.")
doc.add_paragraph()
add_para("Vigneshwara Chinnadurai", bold=True)
add_para("May 2026")

page_break()

# ============================================================
# TABLE OF CONTENTS
# ============================================================
add_heading_styled("TABLE OF CONTENTS", level=1)
doc.add_paragraph()

toc_entries = [
    ("", "Title Page", "1"),
    ("", "Bonafide Certificate", "2"),
    ("", "Declaration", "3"),
    ("", "Acknowledgments", "4"),
    ("", "Table of Contents", "5-6"),
    ("", "List of Tables", "7"),
    ("", "List of Figures", "8"),
    ("", "List of Abbreviations", "9-10"),
    ("", "Executive Summary", "11-12"),
    ("1", "Introduction", "13"),
    ("1.1", "Background of the Study", "13"),
    ("1.2", "Statement of the Problem", "16"),
    ("1.3", "Research Objectives", "18"),
    ("1.4", "Research Questions", "19"),
    ("1.5", "Scope of the Study", "20"),
    ("1.6", "Significance of the Study", "21"),
    ("2", "Literature Review", "22"),
    ("2.1", "Evolution of Large Language Models", "22"),
    ("2.2", "Fundamentals of Quantum Computing", "25"),
    ("2.3", "Quantum Natural Language Processing", "28"),
    ("2.4", "Hybrid Quantum-Classical Architectures", "31"),
    ("2.5", "Comparative Analysis of Approaches", "34"),
    ("2.6", "Research Gaps", "36"),
    ("3", "Research Methodology", "37"),
    ("3.1", "Research Design", "37"),
    ("3.2", "Data Collection Methods", "38"),
    ("3.3", "Experimental Framework", "39"),
    ("3.4", "Tools and Technologies Used", "41"),
    ("3.5", "Data Analysis Techniques", "42"),
    ("3.6", "Ethical Considerations", "43"),
    ("3.7", "Limitations of the Methodology", "44"),
    ("4", "Data Analysis and Interpretation", "45"),
    ("4.1", "Literature Analysis Results", "45"),
    ("4.2", "Experiment 1: Quantum Word Encoding", "48"),
    ("4.3", "Experiment 2: Quantum Text Classification", "51"),
    ("4.4", "Experiment 3: Hybrid Quantum-Classical NLP", "54"),
    ("4.5", "Experiment 4: Performance Benchmarking", "57"),
    ("4.6", "Summary of Experimental Findings", "60"),
    ("5", "Findings and Discussion", "61"),
    ("5.1", "Key Research Findings", "61"),
    ("5.2", "Comparison with Existing Literature", "63"),
    ("5.3", "Practical Implications", "65"),
    ("6", "Conclusions", "66"),
    ("7", "Recommendations", "68"),
    ("7.1", "For Organizations and Industry", "68"),
    ("7.2", "For Academic Research", "69"),
    ("7.3", "For Policy and Education", "70"),
    ("8", "Limitations of the Study", "71"),
    ("9", "References / Bibliography", "73"),
    ("10", "Appendices", "76"),
    ("A", "Appendix A: Experimental Code", "76"),
    ("B", "Appendix B: Full Experimental Results", "82"),
    ("C", "Appendix C: Glossary of Terms", "85"),
]

toc_table = doc.add_table(rows=0, cols=3)
for sec, title, page in toc_entries:
    row = toc_table.add_row()
    row.cells[0].text = sec
    row.cells[1].text = title
    row.cells[2].text = page
    for cell in row.cells:
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11)

page_break()

# ============================================================
# LIST OF TABLES
# ============================================================
add_heading_styled("LIST OF TABLES", level=1)
doc.add_paragraph()

tables_list = [
    ("Table 2.1", "Growth in LLM Parameters and Compute Requirements", "24"),
    ("Table 2.2", "Current Quantum Hardware Landscape (2025)", "27"),
    ("Table 2.3", "Key QNLP Research Timeline", "30"),
    ("Table 2.4", "Comparison of Encoding Strategies", "33"),
    ("Table 2.5", "Literature Classification by Maturity Level", "35"),
    ("Table 3.1", "Experimental Configuration Summary", "40"),
    ("Table 3.2", "Tools and Frameworks Used", "41"),
    ("Table 4.1", "Distribution of Papers by Approach Type", "45"),
    ("Table 4.2", "Distribution of Papers by Research Focus", "46"),
    ("Table 4.3", "Technology Readiness Level Assessment", "47"),
    ("Table 4.4", "Encoding Efficiency Comparison", "49"),
    ("Table 4.5", "Semantic Preservation Results", "49"),
    ("Table 4.6", "Classification Performance Comparison", "52"),
    ("Table 4.7", "Training Convergence Metrics", "53"),
    ("Table 4.8", "Pipeline Comparison Results", "55"),
    ("Table 4.9", "Scalability Analysis", "56"),
    ("Table 4.10", "Noise Impact on Quantum Classifier", "58"),
    ("Table 4.11", "Cross-Framework Comparison", "59"),
    ("Table 4.12", "Summary Comparison Matrix", "60"),
    ("Table 5.1", "Key Findings Summary", "62"),
    ("Table 7.1", "Quantum Readiness Roadmap", "69"),
    ("Table B.1", "Detailed Experiment 1 Results", "82"),
    ("Table B.2", "Detailed Experiment 2 Results", "83"),
    ("Table B.3", "Detailed Experiment 3 Results", "84"),
]

for tnum, title, page in tables_list:
    p = doc.add_paragraph()
    run = p.add_run(f"{tnum}: {title}")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)
    # Add dots and page number
    run2 = p.add_run(f"  {'.' * (60 - len(title))}  {page}")
    run2.font.name = 'Times New Roman'
    run2.font.size = Pt(11)

page_break()

# ============================================================
# LIST OF FIGURES
# ============================================================
add_heading_styled("LIST OF FIGURES", level=1)
doc.add_paragraph()

figures_list = [
    ("Figure 2.1", "Evolution of LLM Capabilities (2017-2026)", "23"),
    ("Figure 2.2", "Quantum Circuit Architecture for NLP", "29"),
    ("Figure 3.1", "Research Methodology Framework", "37"),
    ("Figure 3.2", "Experimental Pipeline Overview", "39"),
    ("Figure 4.1", "Classical Cosine Similarity Heatmap", "48"),
    ("Figure 4.2", "Encoding Comparison: Semantic Preservation", "50"),
    ("Figure 4.3", "Encoding Methods: Comparative Results", "50"),
    ("Figure 4.4", "Quantum Text Classification Results", "53"),
    ("Figure 4.5", "Hybrid Pipeline: Learning Curves", "56"),
    ("Figure 4.6", "Noise Resilience Analysis", "58"),
    ("Figure 4.7", "Evaluation Radar Chart", "59"),
    ("Figure 4.8", "Technology Readiness Assessment", "60"),
    ("Figure 5.1", "Quantum Advantage Zones", "64"),
    ("Figure 7.1", "Phased Adoption Framework", "70"),
]

for fnum, title, page in figures_list:
    p = doc.add_paragraph()
    run = p.add_run(f"{fnum}: {title}")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)
    run2 = p.add_run(f"  {'.' * (60 - len(title))}  {page}")
    run2.font.name = 'Times New Roman'
    run2.font.size = Pt(11)

page_break()

# ============================================================
# LIST OF ABBREVIATIONS
# ============================================================
add_heading_styled("LIST OF ABBREVIATIONS", level=1)
doc.add_paragraph()

abbreviations = [
    ("AI", "Artificial Intelligence"),
    ("API", "Application Programming Interface"),
    ("BERT", "Bidirectional Encoder Representations from Transformers"),
    ("CNOT", "Controlled-NOT (quantum gate)"),
    ("DisCoCat", "Distributional Compositional Categorical"),
    ("DNN", "Deep Neural Network"),
    ("F1", "F1 Score (harmonic mean of precision and recall)"),
    ("GloVe", "Global Vectors for Word Representation"),
    ("GPT", "Generative Pre-trained Transformer"),
    ("GPU", "Graphics Processing Unit"),
    ("IBM", "International Business Machines"),
    ("IQP", "Instantaneous Quantum Polynomial"),
    ("LLM", "Large Language Model"),
    ("ML", "Machine Learning"),
    ("MLP", "Multi-Layer Perceptron"),
    ("MoE", "Mixture of Experts"),
    ("NISQ", "Noisy Intermediate-Scale Quantum"),
    ("NLP", "Natural Language Processing"),
    ("NN", "Neural Network"),
    ("PCA", "Principal Component Analysis"),
    ("PF-days", "Petaflop-days"),
    ("QAOA", "Quantum Approximate Optimization Algorithm"),
    ("QC", "Quantum Computing / Quantum Classifier"),
    ("QML", "Quantum Machine Learning"),
    ("QNLP", "Quantum Natural Language Processing"),
    ("QPU", "Quantum Processing Unit"),
    ("RNN", "Recurrent Neural Network"),
    ("ROI", "Return on Investment"),
    ("SVM", "Support Vector Machine"),
    ("TF-IDF", "Term Frequency-Inverse Document Frequency"),
    ("TPU", "Tensor Processing Unit"),
    ("TRL", "Technology Readiness Level"),
    ("VQC", "Variational Quantum Circuit"),
    ("VQE", "Variational Quantum Eigensolver"),
    ("Word2Vec", "Word to Vector"),
]

abbr_table = doc.add_table(rows=1, cols=2)
abbr_table.style = 'Table Grid'
hdr = abbr_table.rows[0]
hdr.cells[0].text = "Abbreviation"
hdr.cells[1].text = "Full Form"
for cell in hdr.cells:
    for p in cell.paragraphs:
        for run in p.runs:
            run.bold = True
            run.font.name = 'Times New Roman'
    set_cell_shading(cell, "D9E2F3")

for abbr, full in abbreviations:
    row = abbr_table.add_row()
    row.cells[0].text = abbr
    row.cells[1].text = full
    for cell in row.cells:
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11)

page_break()

print("Part 1 complete: Front matter done")
# Save checkpoint
doc.save(OUTPUT_PATH)
print(f"Checkpoint saved: {OUTPUT_PATH}")
