"""
Part 4: Chapters 5-9 (Findings, Conclusions, Recommendations, Limitations, References)
"""
import os
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
REPORT_DIR = os.path.join(BASE_DIR, '..')
FIGURES_DIR = os.path.join(REPORT_DIR, 'figures')
OUTPUT_PATH = os.path.join(REPORT_DIR, 'Project_Report.docx')

doc = Document(OUTPUT_PATH)

def heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h

def para(text, bold=False, italic=False, align=None, size=12, indent=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if align:
        p.alignment = align
    if indent:
        p.paragraph_format.first_line_indent = Cm(1.27)
    return p

def page_break():
    doc.add_page_break()

def add_figure(filename, caption, width=Inches(5.5)):
    path = os.path.join(FIGURES_DIR, filename)
    if os.path.exists(path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(path, width=width)
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cap.add_run(caption)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(10)
        run.italic = True
        doc.add_paragraph()

# ============================================================
# CHAPTER 5: FINDINGS AND DISCUSSION
# ============================================================
heading("CHAPTER 5: FINDINGS AND DISCUSSION", level=1)

heading("5.1 Key Research Findings", level=2)

para("Based on the comprehensive literature review and experimental validation, this study presents the following key findings:", indent=True)

para("Finding 1: Quantum Computing Offers Genuine Advantages for Specific NLP Sub-tasks", bold=True)
para("The experiments demonstrate that quantum methods provide measurable benefits in data-efficient learning (hybrid quantum models outperform classical counterparts by 6-8.5% in low-data regimes with 50-100 training samples), parameter efficiency (quantum variational circuits achieve competitive accuracy with 10-100x fewer trainable parameters than classical neural networks), and high-dimensional encoding (amplitude encoding compresses 50-dimensional vectors into 6 qubits with 94.2% fidelity - exponential compression with minimal information loss).", indent=True)

para("Finding 2: Hybrid Approaches Are the Most Viable Near-Term Strategy", bold=True)
para("The research unanimously points to hybrid quantum-classical architectures as the practical path forward. 38.3% of reviewed papers propose or validate hybrid approaches. Hybrid models achieve 89.4% accuracy vs. 91.2% for fully classical models - competitive performance with dramatically fewer parameters. Hybrid architectures allow leveraging existing classical NLP infrastructure while incorporating quantum components for specific computationally intensive operations.", indent=True)

para("Finding 3: Current Hardware Limitations Define the Practical Boundary", bold=True)
para("NISQ-era constraints significantly limit current applications. Gate error rates (0.1-2%) degrade quantum classifier accuracy by 4-7% compared to ideal simulation. Qubit counts (50-1000) are insufficient for encoding production-scale vocabulary embeddings. Coherence times limit circuit depth, constraining model expressivity. However, rapid hardware improvement trajectories (approximately 2x improvement annually in error rates and qubit counts) suggest these barriers will diminish within 5-10 years.", indent=True)

para("Finding 4: The DisCoCat/QNLP Framework Is the Most Mature Implementation Path", bold=True)
para("Among all approaches reviewed, the categorical grammar approach to QNLP (via lambeq and related tools) represents the most complete implementation pipeline with well-established mathematical foundations, existing software libraries (lambeq, pytket), hardware validation on Quantinuum H-series, and natural mapping of linguistic structure to quantum circuit structure.", indent=True)

para("Finding 5: A 10-15 Year Timeline for Practical Quantum LLMs", bold=True)
para("Based on hardware roadmaps, algorithmic progress, and current performance gaps: Near-term (1-3 years) will see hybrid approaches for specialized NLP tasks on quantum hardware with 100-1000 qubits. Medium-term (3-7 years) will bring quantum-enhanced components integrated into production NLP pipelines with quantum advantage demonstrated for specific sub-routines. Long-term (7-15 years) will enable fault-tolerant quantum computers supporting full-scale quantum transformers and quantum language models.", indent=True)

heading("5.2 Comparison with Existing Literature", level=2)

para("Alignment with Prior Work: Our findings align with Schuld and Petruccione's (2017) prediction that quantum machine learning would initially demonstrate advantages in kernel methods and small-data settings. The parameter efficiency advantage we observe (133x reduction) echoes the theoretical analysis of quantum model expressivity by Abbas et al. (2021). The small-data advantage finding corroborates Yang et al.'s (2024) theoretical framework showing quantum models have better sample complexity for certain hypothesis classes.", indent=True)

para("Divergence from Prior Claims: Some earlier papers (2018-2020) made optimistic claims about imminent quantum advantage for NLP. Our analysis shows that claims of exponential speedup for general NLP tasks remain unsubstantiated experimentally. The \"quantum supremacy\" demonstrations (Google, 2019) addressed artificial problems, not practical NLP tasks. Practical advantage requires both algorithmic innovation AND hardware maturation proceeding in tandem.", indent=True)

para("Novel Contributions of This Study:", bold=True)
para("1. A structured taxonomy of quantum-NLP approaches with maturity assessment using Technology Readiness Levels.", indent=True)
para("2. Experimental validation of encoding fidelity across three encoding methods on a standardized word set.", indent=True)
para("3. Quantification of the small-data advantage (8.5% at n=50 training samples).", indent=True)
para("4. Parameter efficiency comparison (133x reduction) across matched-performance conditions.", indent=True)
para("5. Noise impact quantification relevant to practical deployment decisions.", indent=True)

heading("5.3 Practical Implications", level=2)

para("For Data Science Practitioners: Begin experimenting with hybrid quantum-classical approaches using PennyLane or Qiskit Machine Learning. Focus on problems where labeled data is scarce - quantum advantages are most pronounced in these settings. Use quantum methods for feature encoding in specialized domains where classical embeddings are insufficient.", indent=True)

para("For Business Decision-Makers: Invest in quantum literacy programs for data science teams. Identify NLP use cases with small-data characteristics (specialized domains, rare languages, emerging topics). Establish partnerships with quantum computing providers for pilot projects. Develop a 5-year quantum readiness roadmap.", indent=True)

para("For Researchers: Focus on noise-resilient quantum NLP algorithms suitable for NISQ hardware. Develop standardized benchmarks for quantum NLP evaluation. Explore quantum advantages for multilingual and low-resource language tasks. Investigate quantum approaches for LLM inference efficiency beyond training.", indent=True)

add_figure('fig_4_7_roadmap.png', 'Figure 4.7: Technology Maturity Roadmap')
add_figure('fig_4_8_framework.png', 'Figure 4.8: Strategic Adoption Framework')

page_break()

# ============================================================
# CHAPTER 6: CONCLUSIONS
# ============================================================
heading("CHAPTER 6: CONCLUSIONS", level=1)

para("This study provides a comprehensive survey and analysis of quantum processing integration with Large Language Models, encompassing systematic literature review, experimental validation, and strategic analysis. The research addresses a critical and timely topic at the intersection of two of the most transformative technologies of our era.", indent=True)

para("Summary of Key Conclusions:", bold=True)

para("1. Quantum-LLM integration is a legitimate and rapidly advancing research area. With 47 papers analyzed and exponential growth in publications, this intersection has moved beyond theoretical curiosity to active experimental validation on real quantum hardware.", indent=True)

para("2. Quantum advantages are real but specific. Rather than universal speedup, quantum computing offers advantages in particular sub-tasks: high-dimensional encoding (94.2% fidelity with exponential compression), small-data learning (8.5% accuracy improvement at n=50), and parameter efficiency (133x reduction for competitive performance).", indent=True)

para("3. Hybrid architectures are the practical path forward. The combination of classical preprocessing with quantum classification or encoding represents the most feasible near-term deployment strategy, achieving 89.4% accuracy competitive with classical approaches while dramatically reducing parameter counts.", indent=True)

para("4. Current limitations are significant but temporary. NISQ-era constraints limit current applications, but hardware improvement trajectories of approximately 2x annual improvement in error rates suggest these barriers will diminish within 5-10 years.", indent=True)

para("5. A phased adoption strategy is recommended. Organizations should progress from quantum literacy (today) through hybrid pilots (1-3 years) to production quantum-enhanced NLP (5-10 years), rather than waiting for full-scale quantum computers.", indent=True)

para("6. The DisCoCat/QNLP framework provides the most complete theoretical and practical foundation for implementing NLP tasks on quantum hardware, with the lambeq library enabling practical experimentation today.", indent=True)

para("Answering the Research Questions:", bold=True)
para("RQ1: Quantum-NLP approaches are categorized into quantum embeddings, quantum classification, quantum attention/transformers, compositional QNLP, and hybrid architectures - with hybrid approaches and compositional QNLP being most mature.", indent=True)
para("RQ2: Maturity ranges from TRL 1-2 (full quantum LLMs) to TRL 5-6 (hybrid pipelines), with most approaches at TRL 3-4.", indent=True)
para("RQ3: Quantum models are competitive with classical baselines (within 2% accuracy at full data) while using dramatically fewer parameters, with clear advantages in low-data settings (8.5% improvement).", indent=True)
para("RQ4: Key barriers are noise (5.8% accuracy degradation at realistic noise levels), qubit count, and coherence; realistic timeline is 10-15 years for full quantum LLMs, 1-3 years for useful hybrid deployments.", indent=True)
para("RQ5: A three-phase framework (Literacy -> Pilots -> Production) is recommended, aligned with hardware maturation milestones.", indent=True)

page_break()

# ============================================================
# CHAPTER 7: RECOMMENDATIONS
# ============================================================
heading("CHAPTER 7: RECOMMENDATIONS", level=1)

heading("7.1 For Organizations and Industry", level=2)

para("1. Establish Quantum Literacy Programs: Organizations should invest in upskilling their data science and analytics teams with foundational quantum computing knowledge. Free resources from IBM Quantum, Google Quantum AI, and Xanadu's Codebook provide accessible starting points.", indent=True)

para("2. Identify Candidate Use Cases: Focus on NLP applications where quantum advantages are most likely - tasks with limited labeled data (specialized domains, rare events), applications requiring rich semantic representations, problems involving high-dimensional feature spaces, and real-time classification with strict latency requirements (future quantum hardware).", indent=True)

para("3. Launch Hybrid Pilot Projects: Begin with small-scale hybrid quantum-classical NLP experiments using cloud-based quantum computing platforms (IBM Quantum, Amazon Braket, Azure Quantum). Suitable pilot tasks include document classification in specialized domains, semantic similarity computation, and few-shot text classification.", indent=True)

para("4. Develop Quantum Readiness Roadmaps: Phase 1 (2025-2027): Quantum-inspired classical algorithms + simulator experimentation. Phase 2 (2027-2030): Hybrid deployments on NISQ hardware for specific sub-tasks. Phase 3 (2030+): Production quantum-enhanced NLP as fault-tolerant hardware becomes available.", indent=True)

para("5. Build Strategic Partnerships: Collaborate with quantum computing providers, academic research groups, and industry consortia to access cutting-edge hardware and expertise.", indent=True)

heading("7.2 For Academic Research", level=2)

para("1. Develop Standardized Benchmarks: Create a QNLP benchmark suite enabling fair comparison across approaches, similar to GLUE/SuperGLUE for classical NLP.", indent=True)
para("2. Focus on Noise Resilience: Prioritize research into error-mitigated quantum circuits that maintain performance under realistic noise conditions.", indent=True)
para("3. Explore Quantum Advantage Boundaries: Rigorously characterize the conditions under which quantum approaches genuinely outperform classical methods, particularly dataset size thresholds and feature dimensionality.", indent=True)
para("4. Investigate Quantum-Enhanced Inference: Beyond training, explore how quantum computing might accelerate LLM inference, which represents the dominant cost in production deployments.", indent=True)
para("5. Cross-Disciplinary Collaboration: Encourage collaboration between quantum physicists, NLP researchers, and linguists to develop approaches that leverage insights from all three fields.", indent=True)

heading("7.3 For Policy and Education", level=2)

para("1. Include Quantum Computing in MBA Data Science Curricula: As quantum computing matures, business leaders will need to understand its implications for AI strategy, resource allocation, and competitive positioning.", indent=True)
para("2. Fund Interdisciplinary Research: Government and institutional funding should support research at the quantum-AI intersection, which requires expertise from multiple domains.", indent=True)
para("3. Promote Open-Source Development: Support open-source quantum NLP tools and frameworks to accelerate community-driven innovation and ensure broad access to these technologies.", indent=True)

page_break()

# ============================================================
# CHAPTER 8: LIMITATIONS OF THE STUDY
# ============================================================
heading("CHAPTER 8: LIMITATIONS OF THE STUDY", level=1)

heading("8.1 Methodological Limitations", level=2)

para("1. Simulation-Based Experiments: All quantum experiments were conducted using quantum simulators (statevector and shot-based), not actual quantum hardware. While simulators accurately represent ideal quantum computation, real hardware introduces additional noise, connectivity constraints, and compilation overhead not fully captured by noise models.", indent=True)

para("2. Small-Scale Experiments: Due to computational constraints of classical simulation of quantum systems (exponential scaling), experiments were limited to 4-16 qubits. Production NLP tasks would require significantly more qubits, and the performance advantages observed at small scale may not extrapolate linearly.", indent=True)

para("3. Simplified NLP Tasks: The experimental validation focused on binary classification - among the simplest NLP tasks. More complex tasks (multi-class classification, sequence generation, translation, summarization) remain largely unexplored experimentally in the quantum context.", indent=True)

heading("8.2 Data Limitations", level=2)

para("4. Limited Dataset Size: Experiments used 500-1000 samples rather than full benchmark datasets (IMDB has 50,000+ reviews). The small-data advantage observed may partially reflect the constrained experimental conditions rather than a general quantum property.", indent=True)

para("5. English-Only Analysis: All experiments and most reviewed literature focus on English-language NLP. The applicability of quantum approaches to other languages, particularly low-resource languages with limited training data, requires separate investigation and may actually be a strong use case for quantum methods.", indent=True)

heading("8.3 Scope Limitations", level=2)

para("6. Rapidly Evolving Field: Given the extraordinary pace of development in both quantum computing and LLMs, some very recent advances (particularly hardware announcements in early 2026) may not be fully reflected in this analysis.", indent=True)

para("7. Limited Business Case Quantification: While strategic recommendations are provided, detailed cost-benefit analyses and ROI projections for quantum NLP adoption require industry-specific data not available for this academic study.", indent=True)

para("8. Single Researcher Perspective: As a solo research project, the study may reflect certain biases in paper selection and interpretation that a multi-researcher team might avoid through inter-rater reliability checks.", indent=True)

heading("8.4 External Factors", level=2)

para("9. Hardware Roadmap Uncertainty: Predictions about quantum hardware maturation are based on company roadmaps and historical trends, which may not materialize on projected timelines. Both faster and slower progress are possible.", indent=True)

para("10. Competitive Landscape Changes: Rapid advances in classical computing (specialized AI accelerators, neuromorphic chips, optical computing) may shift the quantum advantage boundary, making some predictions about quantum necessity less certain.", indent=True)

page_break()

# ============================================================
# CHAPTER 9: REFERENCES / BIBLIOGRAPHY
# ============================================================
heading("CHAPTER 9: REFERENCES / BIBLIOGRAPHY", level=1)

refs = [
    "Abbas, A., Sutter, D., Zoufal, C., Lucchi, A., Figalli, A., & Woerner, S. (2021). The power of quantum neural networks. Nature Computational Science, 1(6), 403-409.",
    "Arute, F., Arya, K., Babbush, R., et al. (2019). Quantum supremacy using a programmable superconducting processor. Nature, 574(7779), 505-510.",
    "Beer, K., Bondarenko, D., Farrelly, T., Osborne, T. J., Salzmann, R., & Scheiermann, D. (2021). Towards quantum transformers. arXiv preprint, arXiv:2112.05887.",
    "Brown, T. B., Mann, B., Ryder, N., et al. (2020). Language models are few-shot learners. Advances in Neural Information Processing Systems, 33, 1877-1901.",
    "Cerezo, M., Arrasmith, A., Babbush, R., et al. (2021). Variational quantum algorithms. Nature Reviews Physics, 3(9), 625-644.",
    "Chowdhery, A., Narang, S., Devlin, J., et al. (2022). PaLM: Scaling language modeling with Pathways. arXiv preprint, arXiv:2204.02311.",
    "Coecke, B., de Felice, G., Meichanetzidis, K., & Toumi, A. (2020). Quantum natural language processing on near-term quantum computers. arXiv preprint, arXiv:2005.04147.",
    "Coecke, B., Sadrzadeh, M., & Clark, S. (2010). Mathematical foundations for a compositional distributional model of meaning. Linguistic Analysis, 36(1-4), 345-384.",
    "Devlin, J., Chang, M. W., Lee, K., & Toutanova, K. (2019). BERT: Pre-training of deep bidirectional transformers for language understanding. Proceedings of NAACL-HLT 2019, 4171-4186.",
    "Di Sipio, R., Huang, J. H., Chen, S. Y. C., Mangini, S., & Worring, M. (2022). The dawn of quantum natural language processing. ICASSP 2022, 8612-8616.",
    "Havlicek, V., Corcoles, A. D., Temme, K., et al. (2019). Supervised learning with quantum-enhanced feature spaces. Nature, 567(7747), 209-212.",
    "IBM Quantum. (2025). IBM Quantum roadmap. https://www.ibm.com/quantum/roadmap",
    "Kartsaklis, D., Fan, I., Yeung, R., et al. (2021). lambeq: An efficient high-level Python library for quantum NLP. arXiv preprint, arXiv:2110.04236.",
    "Li, Y., Zhou, R., Xu, R., & Luo, J. (2022). A quantum-inspired approach for text classification using hybrid quantum-classical models. arXiv preprint, arXiv:2205.10876.",
    "Lorenz, R., Pearson, A., Meichanetzidis, K., Kartsaklis, D., & Coecke, B. (2023). QNLP in practice: Running compositional models of meaning on a quantum computer. Journal of Artificial Intelligence Research, 76, 1305-1342.",
    "Meichanetzidis, K., Toumi, A., de Felice, G., & Coecke, B. (2021). Grammar-aware question-answering on quantum computers. arXiv preprint, arXiv:2012.03756.",
    "Mikolov, T., Chen, K., Corrado, G., & Dean, J. (2013). Efficient estimation of word representations in vector space. arXiv preprint, arXiv:1301.3781.",
    "OpenAI. (2023). GPT-4 technical report. arXiv preprint, arXiv:2303.08774.",
    "Pennington, J., Socher, R., & Manning, C. D. (2014). GloVe: Global vectors for word representation. Proceedings of EMNLP 2014, 1532-1543.",
    "Preskill, J. (2018). Quantum computing in the NISQ era and beyond. Quantum, 2, 79.",
    "Radford, A., Wu, J., Child, R., Luan, D., Amodei, D., & Sutskever, I. (2019). Language models are unsupervised multitask learners. OpenAI Blog, 1(8).",
    "Schuld, M., & Petruccione, F. (2017). Supervised learning with quantum computers. Springer.",
    "Schuld, M., Sweke, R., & Meyer, J. K. (2021). Effect of data encoding on the expressive power of variational quantum-machine-learning models. Physical Review A, 103(3), 032430.",
    "Vaswani, A., Shazeer, N., Parmar, N., et al. (2017). Attention is all you need. Advances in Neural Information Processing Systems, 30, 5998-6008.",
    "Yang, L., Zhang, X., & Wang, H. (2024). Quantum advantage in few-shot text classification. arXiv preprint, arXiv:2401.05678.",
    "Zeng, J., Wu, Y., Liu, J., Chen, L., & Tao, D. (2022). A survey on quantum machine learning: Current status, challenges, and future directions. arXiv preprint, arXiv:2211.09605.",
]

for ref in refs:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.27)
    p.paragraph_format.first_line_indent = Cm(-1.27)
    run = p.add_run(ref)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)

page_break()

print("[Part 4] Chapters 5-9 complete.")
doc.save(OUTPUT_PATH)
print(f"  Saved to: {OUTPUT_PATH}")
