"""
Tools for the Resume Parser agent.

The actual LLM-driven extraction happens inside the agent via
`Agent.structured_output(Resume, raw_text)` — these tools just handle I/O.
"""
from __future__ import annotations

import json
from pathlib import Path

from strands import tool

from ..config import DATA_DIR


@tool
def read_resume_file(path: str) -> str:
    """
    Read a resume file (PDF, DOCX, or plain text) and return its text content.

    Args:
        path: Path to the resume file. Relative paths are resolved against the
              project root.

    Returns:
        The extracted plain-text content of the resume.
    """
    file_path = Path(path)
    if not file_path.is_absolute():
        file_path = (DATA_DIR.parent / path).resolve()

    if not file_path.exists():
        return f"ERROR: file not found at {file_path}"

    suffix = file_path.suffix.lower()

    try:
        if suffix == ".pdf":
            from pypdf import PdfReader

            reader = PdfReader(str(file_path))
            return "\n\n".join(
                (page.extract_text() or "") for page in reader.pages
            )

        if suffix in {".docx"}:
            from docx import Document

            doc = Document(str(file_path))
            parts = [p.text for p in doc.paragraphs if p.text.strip()]
            # Also grab table cells (resumes often use tables for layout)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            parts.append(cell.text)
            return "\n".join(parts)

        if suffix in {".txt", ".md"}:
            return file_path.read_text(encoding="utf-8", errors="ignore")

        return f"ERROR: unsupported file type {suffix}. Use PDF, DOCX, or TXT."

    except Exception as e:
        return f"ERROR reading resume: {e!r}"


@tool
def save_profile(profile_json: str) -> str:
    """
    Persist the parsed resume profile to data/profile.json so we don't have to
    re-parse on every run.

    Args:
        profile_json: The parsed profile as a JSON string.

    Returns:
        The path where the profile was saved.
    """
    target = DATA_DIR / "profile.json"
    target.parent.mkdir(parents=True, exist_ok=True)

    # Validate it's real JSON before writing
    try:
        parsed = json.loads(profile_json)
    except json.JSONDecodeError as e:
        return f"ERROR: invalid JSON — {e}"

    target.write_text(
        json.dumps(parsed, indent=2, ensure_ascii=False), encoding="utf-8"
    )
    return f"Saved profile to {target}"


@tool
def load_profile() -> str:
    """
    Load the cached profile from data/profile.json, if it exists.

    Returns:
        The profile JSON as a string, or an empty string if no cache exists.
    """
    target = DATA_DIR / "profile.json"
    if not target.exists():
        return ""
    return target.read_text(encoding="utf-8")
