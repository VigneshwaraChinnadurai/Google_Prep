from pathlib import Path
import re
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


BASE = Path(r"c:/Users/vichinnadurai/Documents/Vignesh/Personal Enrichment/Personal Work/Google_Prep/ML_Interview_Prep")
OUT = BASE / "printable_docx_topics"
OUT.mkdir(exist_ok=True)
PRINT_OUT = BASE / "printable_docx_topics_print_friendly"
PRINT_OUT.mkdir(exist_ok=True)

ALL_MD = sorted([p.name for p in BASE.glob("*.md")])

TOPIC_MAP = {
    "01_Traditional_ML_Interview_Guide.docx": [
        "01_traditional_ml_comprehensive.md",
        "traditional_ml_concepts.md",
    ],
    "02_Computer_Vision_Interview_Guide.docx": [
        "02_computer_vision_deep_learning.md",
        "computer_vision_deep_learning.md",
    ],
    "03_NLP_Interview_Guide.docx": [
        "03_nlp_deep_learning.md",
        "nlp_deep_learning.md",
    ],
    "04_Forecasting_Interview_Guide.docx": [
        "04_forecasting_deep_learning.md",
        "forecasting_deep_learning.md",
    ],
    "05_Recommendation_Systems_Interview_Guide.docx": [
        "05_recommendation_engine_deep_learning.md",
        "recommendation_engines_deep_learning.md",
    ],
    "06_RAG_GenAI_Interview_Guide.docx": ["06_rag_gen_ai.md"],
    "07_AWS_AgentCore_Agentic_AI_Interview_Guide.docx": ["07_agentcore_aws_agentic_ai.md"],
    "08_Quantization_and_LLM_Training_Interview_Guide.docx": ["08_quantization_llm_training.md"],
    "09_Agentic_AI_Safety_Security_Guardrails_Interview_Guide.docx": [
        "09_agentic_hallucination_security_guardrails.md"
    ],
}

COVERED = sorted({f for v in TOPIC_MAP.values() for f in v})
MISSING = [f for f in ALL_MD if f not in COVERED]
if MISSING:
    raise SystemExit(f"Unmapped markdown files: {MISSING}")

INLINE_PAT = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`)")


def add_inline_runs(paragraph, text: str) -> None:
    parts = INLINE_PAT.split(text)
    for part in parts:
        if not part:
            continue
        run = paragraph.add_run()
        if part.startswith("**") and part.endswith("**") and len(part) >= 4:
            run.text = part[2:-2]
            run.bold = True
        elif part.startswith("`") and part.endswith("`") and len(part) >= 2:
            run.text = part[1:-1]
            run.font.name = "Consolas"
        else:
            run.text = part


def is_table_sep(cells) -> bool:
    for cell in cells:
        t = cell.strip().replace(":", "").replace("-", "")
        if t != "":
            return False
    return True


def parse_table_block(doc: Document, table_lines) -> None:
    rows = []
    for ln in table_lines:
        t = ln.strip().strip("|")
        cells = [c.strip() for c in t.split("|")]
        rows.append(cells)

    if not rows:
        return

    clean = []
    for i, row in enumerate(rows):
        if i == 1 and is_table_sep(row):
            continue
        clean.append(row)

    if not clean:
        return

    col_count = max(len(r) for r in clean)
    table = doc.add_table(rows=len(clean), cols=col_count)
    table.style = "Table Grid"

    for i, row in enumerate(clean):
        for j in range(col_count):
            val = row[j] if j < len(row) else ""
            table.cell(i, j).text = val


def configure_print_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    title = doc.styles["Title"]
    title.font.name = "Calibri"
    title.font.size = Pt(20)
    title.font.bold = True

    for heading_name, size in [
        ("Heading 1", 16),
        ("Heading 2", 14),
        ("Heading 3", 12),
        ("Heading 4", 11),
    ]:
        style = doc.styles[heading_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True


def add_markdown(doc: Document, md: str) -> None:
    lines = md.splitlines()
    in_code = False
    table_buf = []

    for raw in lines:
        line = raw.rstrip("\n")
        s = line.strip()

        if s.startswith("```"):
            if table_buf:
                parse_table_block(doc, table_buf)
                table_buf = []
            in_code = not in_code
            continue

        if in_code:
            p = doc.add_paragraph(style="No Spacing")
            r = p.add_run(line)
            r.font.name = "Consolas"
            r.font.size = Pt(9)
            continue

        if "|" in line and line.count("|") >= 2:
            table_buf.append(line)
            continue
        elif table_buf:
            parse_table_block(doc, table_buf)
            table_buf = []

        if s == "":
            doc.add_paragraph("")
            continue

        if re.match(r"^#{1,6}\s+", s):
            level = len(s) - len(s.lstrip("#"))
            txt = s[level:].strip()
            style = f"Heading {min(level, 4)}" if level > 1 else "Heading 1"
            p = doc.add_paragraph(style=style)
            add_inline_runs(p, txt)
            continue

        if re.match(r"^[-*+]\s+", s):
            txt = re.sub(r"^[-*+]\s+", "", s)
            p = doc.add_paragraph(style="List Bullet")
            add_inline_runs(p, txt)
            continue

        if re.match(r"^\d+\.\s+", s):
            txt = re.sub(r"^\d+\.\s+", "", s)
            p = doc.add_paragraph(style="List Number")
            add_inline_runs(p, txt)
            continue

        if re.match(r"^---+$", s):
            continue

        p = doc.add_paragraph()
        add_inline_runs(p, line)

    if table_buf:
        parse_table_block(doc, table_buf)


for out_name, files in TOPIC_MAP.items():
    doc = Document()
    configure_print_styles(doc)
    sec = doc.sections[0]
    sec.top_margin = Inches(0.6)
    sec.bottom_margin = Inches(0.6)
    sec.left_margin = Inches(0.6)
    sec.right_margin = Inches(0.6)

    title = out_name.replace("_", " ").replace(".docx", "")
    p = doc.add_paragraph(style="Title")
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p.add_run(title)

    meta = doc.add_paragraph()
    meta.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    meta_run = meta.add_run(
        "Compiled from source topics in ML_Interview_Prep for quick interview revision"
    )
    meta_run.italic = True

    doc.add_page_break()

    doc.add_paragraph("Source files included:", style="Heading 2")
    for f in files:
        doc.add_paragraph(f, style="List Bullet")

    doc.add_page_break()

    for i, f in enumerate(files):
        md = (BASE / f).read_text(encoding="utf-8", errors="ignore").replace("\ufeff", "")
        doc.add_paragraph(f"Section from {f}", style="Heading 1")
        add_markdown(doc, md)
        if i < len(files) - 1:
            doc.add_page_break()

    target = OUT / out_name
    doc.save(target)
    print(f"Created: {target}")

    print_target = PRINT_OUT / out_name
    doc.save(print_target)
    print(f"Created: {print_target}")

print(f"\nDone. Generated {len(TOPIC_MAP)} DOCX files in {OUT}")
print(f"Done. Generated {len(TOPIC_MAP)} DOCX files in {PRINT_OUT}")
